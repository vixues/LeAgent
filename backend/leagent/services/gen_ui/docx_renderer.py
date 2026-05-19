"""Render a validated GenUI tree to a Word document (.docx)."""

from __future__ import annotations

import io
from typing import Any


def render_genui_to_docx(normalized_tree: dict[str, Any], *, mode: str) -> bytes:
    """Convert a normalized GenUI tree into a .docx file and return its bytes."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()

    root = normalized_tree.get("root")
    if not isinstance(root, dict):
        doc.add_paragraph("(empty)")
        return _save_to_bytes(doc)

    root_kind = str(root.get("kind") or "")

    if mode == "deck" and root_kind == "SlideDeck":
        slides = [
            c for c in (root.get("children") or [])
            if isinstance(c, dict) and str(c.get("kind")) == "Slide"
        ]
        if not slides:
            slides = [c for c in (root.get("children") or []) if isinstance(c, dict)]
        for i, slide in enumerate(slides):
            if i > 0:
                doc.add_page_break()
            _render_slide_to_docx(doc, slide, Inches, Pt, WD_ALIGN_PARAGRAPH)
    else:
        _render_node_to_docx(doc, root, Inches, Pt, WD_ALIGN_PARAGRAPH)

    return _save_to_bytes(doc)


def _save_to_bytes(doc: Any) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_slide_to_docx(doc: Any, node: dict[str, Any], Inches: Any, Pt: Any, WD_ALIGN: Any) -> None:
    props = node.get("props") or {}
    title = props.get("title", "")
    subtitle = props.get("subtitle", "")
    eyebrow = props.get("eyebrow", "")

    if eyebrow:
        p = doc.add_paragraph()
        run = p.add_run(eyebrow.upper())
        run.font.size = Pt(9)
        run.bold = True

    if title:
        doc.add_heading(title, level=1)

    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(12)

    for child in node.get("children") or []:
        if isinstance(child, dict):
            _render_node_to_docx(doc, child, Inches, Pt, WD_ALIGN)


def _render_node_to_docx(doc: Any, node: dict[str, Any], Inches: Any, Pt: Any, WD_ALIGN: Any) -> None:
    kind = str(node.get("kind") or "")
    props: dict[str, Any] = node.get("props") or {}
    children = [c for c in (node.get("children") or []) if isinstance(c, dict)]

    if kind == "Text":
        doc.add_paragraph(str(props.get("value", "")))

    elif kind == "Heading":
        level = min(4, max(1, int(props.get("level") or 2)))
        doc.add_heading(str(props.get("value", "")), level=level)

    elif kind == "Markdown":
        doc.add_paragraph(str(props.get("content", "")))

    elif kind == "Divider":
        doc.add_paragraph("─" * 40)

    elif kind == "Image":
        src = props.get("src", "")
        if src:
            doc.add_paragraph(f"[Image: {src}]")

    elif kind == "Table":
        headers = props.get("headers") or []
        rows_data: list[list[str]] = []
        if headers:
            rows_data.append([str(h) for h in headers])
        for child in children:
            if str(child.get("kind")) == "TableRow":
                row_cells = []
                for cell in (child.get("children") or []):
                    if isinstance(cell, dict):
                        cell_props = cell.get("props") or {}
                        row_cells.append(str(cell_props.get("value", "")))
                rows_data.append(row_cells)

        if rows_data:
            num_cols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.style = "Table Grid"
            for i, row_data in enumerate(rows_data):
                for j, cell_text in enumerate(row_data):
                    if j < num_cols:
                        table.cell(i, j).text = cell_text
            if headers:
                for cell in table.rows[0].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    elif kind == "SectionHeader":
        eyebrow = props.get("eyebrow", "")
        title = props.get("title", "")
        description = props.get("description", "")
        if eyebrow:
            p = doc.add_paragraph()
            run = p.add_run(eyebrow.upper())
            run.font.size = Pt(9)
            run.bold = True
        if title:
            doc.add_heading(title, level=2)
        if description:
            doc.add_paragraph(description)

    elif kind == "KeyValueList":
        items = props.get("items") or []
        for item in items:
            if isinstance(item, dict):
                label = str(item.get("label", ""))
                value = str(item.get("value", ""))
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.bold = True
                p.add_run(value)

    elif kind == "QuoteCard":
        quote = props.get("quote", "")
        author = props.get("author", "")
        p = doc.add_paragraph(style="Quote")
        p.add_run(f'"{quote}"')
        if author:
            p2 = doc.add_paragraph()
            p2.add_run(f"— {author}").italic = True

    elif kind in ("Stack", "Row", "Grid", "ScrollArea", "Card", "DesignSurface", "AspectBox"):
        for child in children:
            _render_node_to_docx(doc, child, Inches, Pt, WD_ALIGN)

    elif kind == "Stat":
        value = props.get("value", "")
        label = props.get("label", "")
        p = doc.add_paragraph()
        run = p.add_run(str(value))
        run.bold = True
        run.font.size = Pt(16)
        if label:
            p.add_run(f"  {label}")

    elif kind == "List" or kind == "ListItem":
        if kind == "List":
            for child in children:
                _render_node_to_docx(doc, child, Inches, Pt, WD_ALIGN)
        else:
            doc.add_paragraph(str(props.get("value", props.get("text", ""))), style="List Bullet")

    elif kind == "CodeBlock":
        code = props.get("code", props.get("value", ""))
        p = doc.add_paragraph()
        run = p.add_run(str(code))
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    elif kind == "FeatureGrid":
        items = props.get("items") or []
        for item in items:
            if isinstance(item, dict):
                title = str(item.get("title", ""))
                desc = str(item.get("description", ""))
                p = doc.add_paragraph()
                run = p.add_run(title)
                run.bold = True
                if desc:
                    p.add_run(f" — {desc}")

    elif kind == "Stepper":
        steps = props.get("steps") or []
        for i, step in enumerate(steps, 1):
            if isinstance(step, dict):
                title = str(step.get("title", ""))
                desc = str(step.get("description", ""))
                doc.add_paragraph(f"{i}. {title}: {desc}", style="List Number")

    else:
        for child in children:
            _render_node_to_docx(doc, child, Inches, Pt, WD_ALIGN)
        if not children and props:
            text = str(props.get("value") or props.get("label") or props.get("title") or "")
            if text:
                doc.add_paragraph(text)
