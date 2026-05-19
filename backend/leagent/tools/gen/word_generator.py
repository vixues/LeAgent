"""Word Generator Tool - Create Word documents programmatically.

Uses python-docx for creating .docx files with rich formatting.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import structlog

from leagent.tools.base import SyncTool, ToolCategory, ToolContext

logger = structlog.get_logger(__name__)


def _add_toc_field(doc: Any) -> None:
    """Inject a Table of Contents field code into the document.

    python-docx has no built-in TOC support so we insert raw XML.
    The TOC updates automatically when the document is opened in Word
    (File -> Print or Ctrl+A -> F9).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_char_separate)

    run4 = paragraph.add_run("Update this table of contents (Ctrl+A, F9)")

    run5 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_char_end)


def _add_page_number_to_footer(section: Any) -> None:
    """Add a PAGE field to the section footer so each page shows its number."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = 1  # center

    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = para.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = para.add_run("1")

    run5 = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def _add_hyperlink(paragraph: Any, url: str, text: str, color: str = "0563C1") -> None:
    """Append an external hyperlink run to *paragraph*."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    rpr.append(r_style)

    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_cell_shading(cell: Any, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


class WordGeneratorTool(SyncTool):
    """Generate Word documents (.docx) with rich content and formatting.

    Features:
    - Create documents from scratch or templates
    - Add headings, paragraphs, tables, lists, hyperlinks
    - Table of contents (auto-generated from headings)
    - Apply styles, text color, background color
    - Tables with column widths, header shading, and alignment
    - Insert images and page breaks
    - Section breaks with orientation changes
    - Page numbers in header/footer
    - Set document properties
    """

    name = "word_generator"
    description = (
        "Generate Word documents (.docx) with support for headings, paragraphs, "
        "tables (with column widths and shading), lists, images, hyperlinks, "
        "table of contents, section breaks, and custom styling. Supports text "
        "and background colors, page numbers, and templates. Ideal for reports, "
        "memos, letters, proposals, and any structured Word document."
    )
    category = ToolCategory.GEN
    version = "2.0.0"
    timeout_sec = 120
    aliases = ["docx_gen", "word_gen", "create_word", "create_docx"]
    search_hint = (
        "Word docx generate create document report memo letter proposal template "
        "headings paragraphs tables lists images hyperlinks TOC page numbers"
    )
    is_concurrency_safe = False
    is_read_only = False
    interrupt_behavior = "cancel"
    max_result_size_chars = 50_000
    path_params = ("template_path",)
    output_path_params = ("output_path",)

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "output_path": {
                    "type": "string",
                    "description": (
                        "Use a bare filename for the generated Word document "
                        "(for example, 'document.docx'); it will be placed in the "
                        "session workspace and shown in the Files tab. Use only "
                        "when the user asked to save or export."
                    ),
                },
                "template_path": {
                    "type": "string",
                    "description": "Optional path to a Word template (.dotx or .docx) to use as base.",
                },
                "title": {
                    "type": "string",
                    "description": "Document title (metadata).",
                },
                "author": {
                    "type": "string",
                    "description": "Document author (metadata).",
                },
                "subject": {
                    "type": "string",
                    "description": "Document subject (metadata).",
                },
                "content": {
                    "type": "array",
                    "description": "Array of content blocks to add to the document.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {
                                "type": "string",
                                "enum": [
                                    "heading",
                                    "paragraph",
                                    "table",
                                    "list",
                                    "image",
                                    "page_break",
                                    "toc",
                                    "hyperlink",
                                    "section_break",
                                ],
                                "description": "Type of content block.",
                            },
                            "text": {
                                "type": "string",
                                "description": "Text content (for heading, paragraph, hyperlink).",
                            },
                            "level": {
                                "type": "integer",
                                "minimum": 1,
                                "maximum": 9,
                                "description": "Heading level (1-9) for heading type.",
                            },
                            "style": {
                                "type": "string",
                                "description": "Style name to apply (e.g., 'Intense Quote', 'List Bullet').",
                            },
                            "bold": {
                                "type": "boolean",
                                "description": "Apply bold formatting.",
                            },
                            "italic": {
                                "type": "boolean",
                                "description": "Apply italic formatting.",
                            },
                            "underline": {
                                "type": "boolean",
                                "description": "Apply underline formatting.",
                            },
                            "font_size": {
                                "type": "integer",
                                "description": "Font size in points.",
                            },
                            "font_name": {
                                "type": "string",
                                "description": (
                                    "Font family name as installed on the machine that opens the "
                                    "document (e.g. 'Arial'). For mixed Chinese + English, pick an "
                                    "installed pan-Unicode face per OS or omit this field to use "
                                    "Word defaults; names differ between Windows, macOS, and Linux."
                                ),
                            },
                            "alignment": {
                                "type": "string",
                                "enum": ["left", "center", "right", "justify"],
                                "description": "Text alignment.",
                            },
                            "text_color": {
                                "type": "string",
                                "description": "Text color as hex code (e.g. 'FF0000' for red).",
                            },
                            "bg_color": {
                                "type": "string",
                                "description": "Paragraph background/shading color as hex code.",
                            },
                            "rows": {
                                "type": "array",
                                "description": "Table rows (for table type).",
                                "items": {
                                    "type": "array",
                                    "items": {"type": "string"},
                                },
                            },
                            "header_row": {
                                "type": "boolean",
                                "description": "Whether first row is a header (for table type).",
                            },
                            "column_widths": {
                                "type": "array",
                                "description": "Column widths in inches (for table type). Must match column count.",
                                "items": {"type": "number"},
                            },
                            "header_shading": {
                                "type": "string",
                                "description": "Header row background color as hex (e.g. 'D5E8F0').",
                            },
                            "items": {
                                "type": "array",
                                "description": "List items (for list type).",
                                "items": {"type": "string"},
                            },
                            "ordered": {
                                "type": "boolean",
                                "description": "Whether list is ordered/numbered (for list type).",
                            },
                            "image_path": {
                                "type": "string",
                                "description": "Path to image file (for image type).",
                            },
                            "width_inches": {
                                "type": "number",
                                "description": "Image width in inches.",
                            },
                            "height_inches": {
                                "type": "number",
                                "description": "Image height in inches.",
                            },
                            "url": {
                                "type": "string",
                                "description": "URL for hyperlink type.",
                            },
                            "orientation": {
                                "type": "string",
                                "enum": ["portrait", "landscape"],
                                "description": "Page orientation after section break.",
                            },
                        },
                        "required": ["type"],
                    },
                },
                "page_setup": {
                    "type": "object",
                    "description": "Page setup options.",
                    "properties": {
                        "orientation": {
                            "type": "string",
                            "enum": ["portrait", "landscape"],
                            "description": "Page orientation.",
                        },
                        "top_margin": {
                            "type": "number",
                            "description": "Top margin in inches.",
                        },
                        "bottom_margin": {
                            "type": "number",
                            "description": "Bottom margin in inches.",
                        },
                        "left_margin": {
                            "type": "number",
                            "description": "Left margin in inches.",
                        },
                        "right_margin": {
                            "type": "number",
                            "description": "Right margin in inches.",
                        },
                    },
                },
                "header_text": {
                    "type": "string",
                    "description": "Text to add to document header.",
                },
                "footer": {
                    "type": "object",
                    "description": "Footer configuration.",
                    "properties": {
                        "text": {
                            "type": "string",
                            "description": "Footer text.",
                        },
                        "include_page_number": {
                            "type": "boolean",
                            "description": "Include automatic page numbers in footer.",
                        },
                    },
                },
                "footer_text": {
                    "type": "string",
                    "description": "Simple footer text (legacy; prefer 'footer' object).",
                },
            },
            "required": ["output_path", "content"],
            "additionalProperties": False,
        }

    def get_activity_description(self, params: dict[str, Any] | None = None) -> str | None:
        return "Generating Word document"

    def execute_sync(self, params: dict[str, Any], context: ToolContext) -> dict[str, Any]:
        """Generate a Word document with the specified content."""
        try:
            from docx import Document
            from docx.enum.section import WD_ORIENT
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt, RGBColor
        except ImportError as e:
            raise RuntimeError(
                "python-docx is not installed. Install with: pip install python-docx"
            ) from e

        output_path = Path(params["output_path"])
        template_path = params.get("template_path")
        content_blocks = params.get("content", [])
        page_setup = params.get("page_setup", {})

        output_path.parent.mkdir(parents=True, exist_ok=True)

        if template_path:
            template = Path(template_path)
            if not template.exists():
                raise FileNotFoundError(f"Template not found: {template}")
            logger.info("Creating document from template", template=str(template))
            doc = Document(str(template))
        else:
            logger.info("Creating new Word document")
            doc = Document()

        if params.get("title"):
            doc.core_properties.title = params["title"]
        if params.get("author"):
            doc.core_properties.author = params["author"]
        if params.get("subject"):
            doc.core_properties.subject = params["subject"]

        if page_setup:
            section = doc.sections[0]
            if page_setup.get("orientation") == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                new_width = section.page_height
                new_height = section.page_width
                section.page_width = new_width
                section.page_height = new_height
            if page_setup.get("top_margin") is not None:
                section.top_margin = Inches(page_setup["top_margin"])
            if page_setup.get("bottom_margin") is not None:
                section.bottom_margin = Inches(page_setup["bottom_margin"])
            if page_setup.get("left_margin") is not None:
                section.left_margin = Inches(page_setup["left_margin"])
            if page_setup.get("right_margin") is not None:
                section.right_margin = Inches(page_setup["right_margin"])

        if params.get("header_text"):
            section = doc.sections[0]
            header = section.header
            header.is_linked_to_previous = False
            header.paragraphs[0].text = params["header_text"]

        footer_config = params.get("footer") or {}
        footer_text = params.get("footer_text", "")
        if footer_config or footer_text:
            section = doc.sections[0]
            ft = footer_config.get("text", "") or footer_text
            if ft:
                footer = section.footer
                footer.is_linked_to_previous = False
                footer.paragraphs[0].text = ft
            if footer_config.get("include_page_number"):
                _add_page_number_to_footer(section)

        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        stats = {
            "headings": 0,
            "paragraphs": 0,
            "tables": 0,
            "lists": 0,
            "images": 0,
            "page_breaks": 0,
            "hyperlinks": 0,
            "toc": 0,
            "section_breaks": 0,
        }

        for block in content_blocks:
            block_type = block.get("type")

            if block_type == "heading":
                level = block.get("level", 1)
                text = block.get("text", "")
                para = doc.add_heading(text, level=level)
                self._apply_formatting(para, block, alignment_map, Pt, RGBColor)
                stats["headings"] += 1

            elif block_type == "paragraph":
                text = block.get("text", "")
                style = block.get("style")
                para = doc.add_paragraph(text, style=style)
                self._apply_formatting(para, block, alignment_map, Pt, RGBColor)
                stats["paragraphs"] += 1

            elif block_type == "table":
                self._add_table(doc, block, Inches, Pt, RGBColor, WD_TABLE_ALIGNMENT)
                stats["tables"] += 1

            elif block_type == "list":
                items = block.get("items", [])
                ordered = block.get("ordered", False)
                style = "List Number" if ordered else "List Bullet"

                for item in items:
                    doc.add_paragraph(item, style=style)

                stats["lists"] += 1

            elif block_type == "image":
                image_path = block.get("image_path")
                if not image_path:
                    continue

                img_file = Path(image_path)
                if not img_file.exists():
                    logger.warning("Image not found, skipping", image_path=image_path)
                    continue

                width = Inches(block["width_inches"]) if block.get("width_inches") else None
                height = Inches(block["height_inches"]) if block.get("height_inches") else None

                doc.add_picture(str(img_file), width=width, height=height)
                stats["images"] += 1

            elif block_type == "page_break":
                doc.add_page_break()
                stats["page_breaks"] += 1

            elif block_type == "toc":
                _add_toc_field(doc)
                stats["toc"] += 1

            elif block_type == "hyperlink":
                url = block.get("url", "")
                text = block.get("text", url)
                if url:
                    para = doc.add_paragraph()
                    _add_hyperlink(para, url, text)
                    stats["hyperlinks"] += 1

            elif block_type == "section_break":
                self._add_section_break(doc, block, Inches, WD_ORIENT)
                stats["section_breaks"] += 1

        doc.save(str(output_path))

        file_size = output_path.stat().st_size

        logger.info(
            "Word document generated successfully",
            output_path=str(output_path),
            file_size=file_size,
            **stats,
        )

        return {
            "success": True,
            "output_path": str(output_path),
            "file_size_bytes": file_size,
            "content_stats": stats,
            "metadata": {
                "title": params.get("title", ""),
                "author": params.get("author", ""),
                "subject": params.get("subject", ""),
            },
        }

    def _add_table(
        self,
        doc: Any,
        block: dict[str, Any],
        Inches: Any,
        Pt: Any,
        RGBColor: Any,
        WD_TABLE_ALIGNMENT: Any,
    ) -> None:
        """Create a table with optional column widths and header shading."""
        rows_data = block.get("rows", [])
        if not rows_data:
            return

        num_rows = len(rows_data)
        num_cols = max(len(row) for row in rows_data) if rows_data else 0

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        col_widths = block.get("column_widths")
        if col_widths and len(col_widths) == num_cols:
            for idx, width_in in enumerate(col_widths):
                for row in table.rows:
                    row.cells[idx].width = Inches(width_in)

        for i, row_data in enumerate(rows_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = str(cell_text)

        header_shading = block.get("header_shading")
        if block.get("header_row") and rows_data:
            for cell in table.rows[0].cells:
                if header_shading:
                    _set_cell_shading(cell, header_shading)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    def _add_section_break(
        self,
        doc: Any,
        block: dict[str, Any],
        Inches: Any,
        WD_ORIENT: Any,
    ) -> None:
        """Add a section break with optional orientation change."""
        from docx.enum.section import WD_SECTION_START

        new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        orientation = block.get("orientation")
        if orientation == "landscape":
            new_section.orientation = WD_ORIENT.LANDSCAPE
            new_width = new_section.page_height
            new_height = new_section.page_width
            new_section.page_width = new_width
            new_section.page_height = new_height
        elif orientation == "portrait":
            new_section.orientation = WD_ORIENT.PORTRAIT
            if new_section.page_width > new_section.page_height:
                new_width = new_section.page_height
                new_height = new_section.page_width
                new_section.page_width = new_width
                new_section.page_height = new_height

    def _apply_formatting(
        self,
        paragraph: Any,
        block: dict[str, Any],
        alignment_map: dict[str, Any],
        Pt: Any,
        RGBColor: Any,
    ) -> None:
        """Apply formatting options to a paragraph."""
        if block.get("alignment"):
            alignment = alignment_map.get(block["alignment"])
            if alignment:
                paragraph.alignment = alignment

        if block.get("bg_color"):
            self._set_paragraph_shading(paragraph, block["bg_color"])

        fmt_keys = ("bold", "italic", "underline", "font_size", "font_name", "text_color")
        if any(block.get(k) for k in fmt_keys):
            for run in paragraph.runs:
                if block.get("bold"):
                    run.bold = True
                if block.get("italic"):
                    run.italic = True
                if block.get("underline"):
                    run.underline = True
                if block.get("font_size"):
                    run.font.size = Pt(block["font_size"])
                if block.get("font_name"):
                    run.font.name = block["font_name"]
                if block.get("text_color"):
                    hex_color = block["text_color"].lstrip("#")
                    try:
                        r = int(hex_color[0:2], 16)
                        g = int(hex_color[2:4], 16)
                        b = int(hex_color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                    except (ValueError, IndexError):
                        pass

    @staticmethod
    def _set_paragraph_shading(paragraph: Any, hex_color: str) -> None:
        """Apply background shading to a paragraph."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color.lstrip("#"))
        shd.set(qn("w:val"), "clear")
        p_pr.append(shd)
