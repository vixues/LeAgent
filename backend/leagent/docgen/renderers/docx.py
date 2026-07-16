"""Professional DOCX renderer (python-docx).

CJK correctness: python-docx cannot embed font binaries, so every run gets an
explicit ``w:eastAsia`` font (from the theme) alongside the ascii face — this
is what stops Word/WPS from silently substituting a Latin-only font for
Chinese text. Latin text keeps the theme's ascii font.

Features: native heading styles (Heading 1–6) restyled from the theme, cover
page, TOC field (refreshes on open with Ctrl+A F9), header/footer with real
PAGE-number fields, zebra tables, code blocks, callouts, metrics rows,
matplotlib charts, quotes, nested/task lists, and hyperlinks.
"""

from __future__ import annotations

import io
import re
from typing import TYPE_CHECKING, Any

import structlog

from leagent.docgen.charts import render_chart_png
from leagent.docgen.checklist import checklist_stats, priority_meta, status_meta
from leagent.docgen.images import resolve_image
from leagent.docgen.markdown import parse_inline
from leagent.docgen.mathtext import latex_lines, latex_to_unicode, render_math_png
from leagent.docgen.model import (
    Block,
    CalloutBlock,
    ChartBlock,
    ChecklistBlock,
    ChecklistItem,
    CodeBlock,
    ColumnsBlock,
    DefinitionListBlock,
    DividerBlock,
    DocumentSpec,
    FootnotesBlock,
    HeadingBlock,
    ImageBlock,
    ListBlock,
    ListItem,
    MathBlock,
    MetricsBlock,
    PageBreakBlock,
    ParagraphBlock,
    QuoteBlock,
    SpacerBlock,
    TableBlock,
    TocBlock,
)
from leagent.docgen.omml import latex_to_omml_element
from leagent.docgen.tables import process_table, resolve_table_style
from leagent.docgen.themes import CALLOUT_COLORS, Theme, get_theme

if TYPE_CHECKING:
    from pathlib import Path

    from leagent.docgen.tables import ProcessedTable, TableStyleSpec

logger = structlog.get_logger(__name__)

_HEX_RE = re.compile(r"^#?([0-9A-Fa-f]{6})$")


def _hex(value: str, fallback: str = "000000") -> str:
    m = _HEX_RE.match(value or "")
    return m.group(1).upper() if m else fallback


def render_docx(spec: DocumentSpec, output_path: Path) -> dict[str, Any]:
    """Render a document spec to a .docx file."""
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.shared import Emu, Pt, RGBColor

    theme = get_theme(spec.theme, kind="document")
    fonts = theme.fonts
    warnings: list[str] = []
    stats = {
        "headings": 0,
        "paragraphs": 0,
        "tables": 0,
        "images": 0,
        "charts": 0,
        "lists": 0,
        "code_blocks": 0,
        "callouts": 0,
        "equations": 0,
    }

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    doc = Document()

    # -- core properties -----------------------------------------------
    props = doc.core_properties
    props.title = spec.title or ""
    props.author = spec.author or ""
    props.subject = spec.subject or ""
    if spec.keywords:
        props.keywords = ", ".join(spec.keywords)

    # -- page setup ------------------------------------------------------
    section = doc.sections[0]
    _apply_page_setup(section, spec, WD_ORIENT, Emu)

    # -- fonts on runs ---------------------------------------------------

    def _style_run(
        run: Any,
        *,
        font_name: str,
        size: float | None = None,
        bold: bool | None = None,
        italic: bool | None = None,
        color: str | None = None,
        strike: bool = False,
    ) -> None:
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:eastAsia"), fonts.east_asia)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if strike:
            run.font.strike = True
        if color:
            run.font.color.rgb = RGBColor.from_string(_hex(color))

    def _add_rich_text(
        paragraph: Any,
        text: str,
        *,
        base_font: str | None = None,
        size: float | None = None,
        bold: bool | None = None,
        color: str | None = None,
    ) -> None:
        for span in parse_inline(text):
            if span.link:
                _add_hyperlink(paragraph, span.text, span.link, qn=qn)
                continue
            if span.math:
                # Prefer a native, editable OMML equation; fall back to a
                # high-DPI image, then to Unicode text.
                omath = latex_to_omml_element(span.text, display=False)
                if omath is not None:
                    paragraph._p.append(omath)
                    continue
                math_size = size or theme.sizes.body
                rendered = render_math_png(
                    span.text, font_size=math_size, color=theme.colors.text, dpi=440
                )
                if rendered is not None:
                    png, _, h_pt, _ = rendered
                    run = paragraph.add_run()
                    run.add_picture(io.BytesIO(png), height=Pt(h_pt))
                    continue
                run = paragraph.add_run(latex_to_unicode(span.text))
                _style_run(
                    run, font_name=base_font or fonts.body, size=size, italic=True
                )
                continue
            run = paragraph.add_run(span.text)
            font_name = fonts.mono if span.code else (base_font or fonts.body)
            _style_run(
                run,
                font_name=font_name,
                size=size,
                bold=True if span.bold else bold,
                italic=True if span.italic else None,
                color=color,
                strike=span.strike,
            )
            if span.sup:
                run.font.superscript = True
            elif span.sub:
                run.font.subscript = True
            if span.code:
                _shade_run(run, theme.colors.surface, qn=qn)

    def _body_paragraph(text: str, *, alignment: str | None = None) -> Any:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(theme.spacing.paragraph_spacing)
        pf.line_spacing = theme.spacing.line_spacing
        if alignment:
            p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        elif spec.justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_rich_text(p, text, size=theme.sizes.body)
        return p

    # -- block renderers -------------------------------------------------

    zh = _looks_chinese(spec)
    nums = {"figure": 0, "table": 0, "equation": 0}
    _caption_words = {
        "table": ("表", "Table"),
        "figure": ("图", "Figure"),
        "equation": ("式", "Equation"),
    }
    body_started = False

    def _caption_text(kind: str, caption: str) -> str:
        if not spec.numbered_figures:
            return caption
        nums[kind] += 1
        word = _caption_words[kind][0 if zh else 1]
        return f"**{word} {nums[kind]}**  {caption}"

    def _add_heading(block: HeadingBlock) -> None:
        if spec.section_pages and block.level == 1 and body_started:
            doc.add_page_break()
        h = doc.add_heading(level=min(block.level, 6))
        # add_heading creates an empty paragraph with the built-in style;
        # fill it with themed runs (built-in styles carry Latin-only fonts).
        for run in list(h.runs):
            run._element.getparent().remove(run._element)
        size = theme.sizes.heading(block.level)
        color = theme.colors.primary if block.level <= 2 else theme.colors.text
        _add_rich_text(
            h,
            block.text,
            base_font=fonts.heading,
            size=size,
            bold=True,
            color=color,
        )
        stats["headings"] += 1

    def _add_table(block: TableBlock) -> None:
        pt = process_table(block, theme=theme)
        if not pt.header and not pt.body:
            return
        ts = resolve_table_style(theme, pt.style)

        all_rows = ([pt.header] if pt.header else []) + pt.body
        table = doc.add_table(rows=len(all_rows), cols=pt.col_count)
        table.style = "Table Grid" if ts.grid else "Normal Table"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        usable = section.page_width - section.left_margin - section.right_margin
        for c_idx, frac in enumerate(pt.width_fractions()):
            col_w = Emu(int(usable * frac))
            for col_cell in table.columns[c_idx].cells:
                col_cell.width = col_w
        body_start = 1 if pt.header else 0
        cell_size = max(theme.sizes.small, theme.sizes.body - 1)

        for r_idx, row in enumerate(all_rows):
            is_header = bool(pt.header) and r_idx == 0
            body_idx = r_idx - body_start
            is_total = not is_header and body_idx == pt.total_row_index
            for c_idx, pcell in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                if pcell.align in ("center", "right"):
                    p.alignment = align_map[pcell.align]
                if is_header:
                    color = ts.header_text if ts.header_fill else theme.colors.primary
                elif pcell.polarity:
                    color = ts.positive if pcell.polarity == "positive" else ts.negative
                else:
                    color = None
                _add_rich_text(
                    p,
                    pcell.text,
                    size=cell_size,
                    bold=True if (is_header or pcell.bold) else None,
                    color=color,
                )
                if is_header and ts.header_fill:
                    _shade_cell(cell, _hex(ts.header_fill), qn=qn)
                elif is_total and ts.total_fill:
                    _shade_cell(cell, _hex(ts.total_fill), qn=qn)
                elif (
                    pt.zebra
                    and ts.zebra_fill
                    and not is_header
                    and body_idx % 2 == 1
                ):
                    _shade_cell(cell, _hex(ts.zebra_fill), qn=qn)

        if not ts.grid:
            _set_table_rules(table, pt, ts, qn=qn)

        if pt.caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_rich_text(
                cap,
                _caption_text("table", pt.caption),
                size=theme.sizes.small,
                color=ts.caption_text,
            )
        stats["tables"] += 1

    def _add_picture_bytes(data: bytes, *, width_pct: float | None, caption: str | None) -> None:
        usable = section.page_width - section.left_margin - section.right_margin
        width = Emu(int(usable * ((width_pct or 88.0) / 100.0)))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(data), width=width)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_rich_text(
                cap,
                _caption_text("figure", caption),
                size=theme.sizes.small,
                color=theme.colors.text_light,
            )

    def _add_image(block: ImageBlock) -> None:
        resolved = resolve_image(
            path=block.path,
            base64_data=block.base64_data,
            url=block.url,
            file_id=block.file_id,
        )
        if resolved is None:
            warnings.append(
                f"Image could not be resolved: "
                f"{block.file_id or block.path or block.url or 'base64'}"
            )
            return
        _add_picture_bytes(resolved.data, width_pct=block.width_pct, caption=block.caption)
        stats["images"] += 1

    def _add_chart(block: ChartBlock) -> None:
        png = render_chart_png(block, theme)
        if png is None:
            warnings.append(f"Chart could not be rendered: {block.title or block.chart_type}")
            return
        _add_picture_bytes(png, width_pct=block.width_pct, caption=block.caption)
        stats["charts"] += 1

    def _add_code(block: CodeBlock) -> None:
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        _shade_cell(cell, _hex(theme.colors.surface), qn=qn)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        lines = block.code.splitlines() or [""]
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break(WD_BREAK.LINE)
            run = p.add_run(line)
            _style_run(run, font_name=fonts.mono, size=theme.sizes.code)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        stats["code_blocks"] += 1

    def _add_callout(block: CalloutBlock) -> None:
        fill_hex, bar_hex = CALLOUT_COLORS.get(block.variant, CALLOUT_COLORS["info"])
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        _shade_cell(cell, _hex(fill_hex), qn=qn)
        _set_cell_left_border(cell, _hex(bar_hex), qn=qn)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if block.title:
            _add_rich_text(
                p, block.title, size=theme.sizes.body, bold=True, color=bar_hex
            )
            if block.text:
                p.add_run().add_break(WD_BREAK.LINE)
        if block.text:
            _add_rich_text(p, block.text, size=theme.sizes.body)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        stats["callouts"] += 1

    def _add_metrics(block: MetricsBlock) -> None:
        if not block.items:
            return
        items = block.items[:5]
        table = doc.add_table(rows=1, cols=len(items))
        table.style = "Table Grid"
        for idx, item in enumerate(items):
            cell = table.cell(0, idx)
            _shade_cell(cell, _hex(theme.colors.surface), qn=qn)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_rich_text(
                p, item.value, size=theme.sizes.h2, bold=True, color=theme.colors.primary
            )
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_rich_text(
                p2, item.label, size=theme.sizes.small, color=theme.colors.text_light
            )
            if item.delta:
                p3 = cell.add_paragraph()
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                delta_color = (
                    "2F9E5B" if not item.delta.strip().startswith("-") else "C0392B"
                )
                _add_rich_text(p3, item.delta, size=theme.sizes.small, color=delta_color)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _add_checklist(block: ChecklistBlock) -> None:
        groups = block.normalized_groups()
        if not groups:
            return
        if block.title:
            p = doc.add_paragraph()
            _add_rich_text(
                p, block.title, size=theme.sizes.h3, bold=True, color=theme.colors.primary
            )
        if block.description:
            p = doc.add_paragraph()
            _add_rich_text(
                p, block.description, size=theme.sizes.body, color=theme.colors.text_light
            )

        stats_data = checklist_stats(block)
        if block.show_progress and stats_data["total_items"]:
            pct = stats_data["progress_percentage"]
            word = "完成度" if zh else "Progress"
            p = doc.add_paragraph()
            _add_rich_text(
                p,
                f"{word}: {pct}% ({stats_data['completed']}/{stats_data['total_items']})",
                size=theme.sizes.small,
                bold=True,
            )
            # Two-cell table acts as a progress bar.
            usable = section.page_width - section.left_margin - section.right_margin
            bar = doc.add_table(rows=1, cols=2)
            bar.autofit = False
            done = max(1, min(99, pct)) if 0 < pct < 100 else max(pct, 1)
            bar.cell(0, 0).width = Emu(int(usable * done / 100))
            bar.cell(0, 1).width = Emu(int(usable * (100 - done) / 100))
            _shade_cell(bar.cell(0, 0), "2F9E5B" if pct > 0 else _hex(theme.colors.surface), qn=qn)
            _shade_cell(bar.cell(0, 1), _hex(theme.colors.surface), qn=qn)
            for cell in bar.rows[0].cells:
                cell.paragraphs[0].add_run(" ")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        def _walk(items: list[ChecklistItem], depth: int) -> None:
            for item in items:
                glyph, color_hex, _ = status_meta(item.status)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(12 + depth * 16)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(glyph + " ")
                _style_run(run, font_name=fonts.body, size=theme.sizes.body, color=color_hex.lstrip("#"), bold=True)
                dim = theme.colors.text_light if item.status in ("completed", "skipped") else None
                _add_rich_text(p, item.text, size=theme.sizes.body, color=dim)
                prio = priority_meta(item.priority)
                if prio:
                    p_color, p_label = prio
                    run = p.add_run(f"  [{p_label.upper()}]")
                    _style_run(run, font_name=fonts.body, size=theme.sizes.small - 1, color=p_color.lstrip("#"), bold=True)
                meta_bits = []
                if item.assignee:
                    meta_bits.append(f"@{item.assignee}")
                if item.due_date:
                    meta_bits.append(f"{'截止' if zh else 'Due'}: {item.due_date}")
                if meta_bits or item.notes:
                    mp = doc.add_paragraph()
                    mp.paragraph_format.left_indent = Pt(24 + depth * 16)
                    mp.paragraph_format.space_after = Pt(1)
                    if meta_bits:
                        _add_rich_text(mp, " · ".join(meta_bits), size=theme.sizes.small, color=theme.colors.text_light)
                    if item.notes:
                        if meta_bits:
                            mp.add_run("  ")
                        _add_rich_text(mp, item.notes, size=theme.sizes.small, color=theme.colors.text_light)
                if item.sub_items:
                    _walk(item.sub_items, depth + 1)

        for group in groups:
            if group.name:
                gp = doc.add_paragraph()
                gp.paragraph_format.space_before = Pt(4)
                _add_rich_text(gp, group.name, size=theme.sizes.body + 1, bold=True)
            if group.description:
                _add_rich_text(
                    doc.add_paragraph(), group.description, size=theme.sizes.small,
                    color=theme.colors.text_light,
                )
            _walk(group.items, 0)

        if block.show_legend:
            legend = doc.add_paragraph()
            legend.paragraph_format.space_before = Pt(4)
            run = legend.add_run("图例: " if zh else "Legend: ")
            _style_run(run, font_name=fonts.body, size=theme.sizes.small, bold=True)
            for status in ("completed", "in_progress", "blocked", "pending"):
                glyph, color_hex, lbl = status_meta(status)
                gr = legend.add_run(f"{glyph} ")
                _style_run(gr, font_name=fonts.body, size=theme.sizes.small, color=color_hex.lstrip("#"), bold=True)
                tr = legend.add_run(f"{lbl}   ")
                _style_run(tr, font_name=fonts.body, size=theme.sizes.small, color=theme.colors.text_light)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _add_list(block: ListBlock) -> None:
        def _walk(items: list[ListItem], depth: int) -> None:
            style = "List Number" if block.ordered and depth == 0 else (
                "List Bullet" if depth == 0 else ("List Bullet 2" if depth == 1 else "List Bullet 3")
            )
            for item in items:
                p = doc.add_paragraph(style=style)
                p.paragraph_format.space_after = Pt(2)
                prefix = ""
                if item.checked is True:
                    prefix = "\u2713 "
                elif item.checked is False:
                    prefix = "\u25a1 "
                if prefix:
                    run = p.add_run(prefix)
                    _style_run(
                        run,
                        font_name=fonts.body,
                        size=theme.sizes.body,
                        color="2F9E5B" if item.checked else None,
                    )
                _add_rich_text(p, item.text, size=theme.sizes.body)
                if item.children:
                    _walk(item.children, depth + 1)

        _walk(block.items, 0)
        stats["lists"] += 1

    def _add_math(block: MathBlock) -> None:
        # Native OMML equations (editable in Word). Multi-line AMS input is
        # stacked as one centered oMathPara per row.
        rows = latex_lines(block.latex)
        omml_rows = [latex_to_omml_element(row, display=True) for row in rows]
        if all(el is not None for el in omml_rows) and omml_rows:
            for el in omml_rows:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p._p.append(el)
            _add_math_caption(block)
            stats["equations"] += 1
            return

        # Fallback: high-DPI image rows, then a code box.
        display_size = theme.sizes.body * 1.18
        rendered_rows: list[tuple[bytes, float]] = []
        for row in rows:
            rendered = render_math_png(
                row, font_size=display_size, color=theme.colors.text, dpi=440
            )
            if rendered is None:
                rendered_rows = []
                break
            png, _, h_pt, _ = rendered
            rendered_rows.append((png, h_pt))
        if not rendered_rows:
            warnings.append(f"Math could not be rendered: {block.latex[:80]}")
            _add_code(CodeBlock(code=block.latex, language="latex"))
            return
        for png, h_pt in rendered_rows:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run()
            run.add_picture(io.BytesIO(png), height=Pt(h_pt))
        _add_math_caption(block)
        stats["equations"] += 1

    def _add_math_caption(block: MathBlock) -> None:
        if not block.caption:
            return
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_rich_text(
            cap,
            _caption_text("equation", block.caption),
            size=theme.sizes.small,
            color=theme.colors.text_light,
        )

    def _add_definition_list(block: DefinitionListBlock) -> None:
        for item in block.items:
            term_p = doc.add_paragraph()
            term_p.paragraph_format.space_after = Pt(1.5)
            _add_rich_text(term_p, item.term, size=theme.sizes.body, bold=True)
            for definition in item.definitions:
                def_p = doc.add_paragraph()
                def_p.paragraph_format.left_indent = Pt(16)
                def_p.paragraph_format.space_after = Pt(3)
                _add_rich_text(def_p, definition, size=theme.sizes.body)

    def _add_footnotes(block: FootnotesBlock) -> None:
        if not block.items:
            return
        _add_divider(doc, theme, qn=qn)
        for item in block.items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            marker = p.add_run(item.label)
            _style_run(
                marker,
                font_name=fonts.body,
                size=theme.sizes.small,
                color=theme.colors.text_light,
            )
            marker.font.superscript = True
            p.add_run(" ")
            _add_rich_text(
                p, item.text, size=theme.sizes.small, color=theme.colors.text_light
            )

    def _add_quote(block: QuoteBlock) -> None:
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        _set_cell_left_border(cell, _hex(theme.colors.primary), qn=qn, size_eighth_pt=24)
        p = cell.paragraphs[0]
        _add_rich_text(p, block.text, size=theme.sizes.body, color=theme.colors.text_light)
        if block.attribution:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_rich_text(
                p2,
                f"— {block.attribution}",
                size=theme.sizes.small,
                color=theme.colors.text_light,
            )
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # -- cover -------------------------------------------------------------

    cover = spec.cover_spec()
    if cover is not None:
        for _ in range(6):
            doc.add_paragraph()
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(cover.title or "Untitled")
        _style_run(
            run,
            font_name=fonts.heading,
            size=theme.sizes.title + 4,
            bold=True,
            color=theme.colors.primary,
        )
        if cover.subtitle:
            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub_p.add_run(cover.subtitle)
            _style_run(
                run,
                font_name=fonts.body,
                size=theme.sizes.h2,
                color=theme.colors.text_light,
            )
        meta_bits = [b for b in (cover.organization, cover.author, cover.date) if b]
        if meta_bits:
            doc.add_paragraph()
            meta_p = doc.add_paragraph()
            meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = meta_p.add_run("  ·  ".join(meta_bits))
            _style_run(
                run,
                font_name=fonts.body,
                size=theme.sizes.body,
                color=theme.colors.text_light,
            )
        doc.add_page_break()
    elif spec.title:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(spec.title)
        _style_run(
            run,
            font_name=fonts.heading,
            size=theme.sizes.title,
            bold=True,
            color=theme.colors.primary,
        )
        if spec.subtitle:
            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub_p.add_run(spec.subtitle)
            _style_run(
                run,
                font_name=fonts.body,
                size=theme.sizes.h3,
                color=theme.colors.text_light,
            )

    # -- header / footer -----------------------------------------------

    _apply_header_footer(section, spec, theme, _style_run, align_map, qn=qn, pt=Pt)

    # -- TOC -------------------------------------------------------------

    toc_blocks = [b for b in spec.blocks if isinstance(b, TocBlock)]
    toc_placed = False

    def _add_toc(title: str | None) -> None:
        head_p = doc.add_paragraph()
        _add_rich_text(
            head_p,
            title or ("目录" if _looks_chinese(spec) else "Contents"),
            base_font=fonts.heading,
            size=theme.sizes.h1,
            bold=True,
            color=theme.colors.primary,
        )
        _insert_toc_field(doc, qn=qn)
        hint = doc.add_paragraph()
        _add_rich_text(
            hint,
            "（在 Word 中按 Ctrl+A 后 F9 刷新目录 / In Word press Ctrl+A then F9 to refresh）",
            size=theme.sizes.small,
            color=theme.colors.text_light,
        )
        doc.add_page_break()

    if spec.toc and not toc_blocks:
        _add_toc(None)
        toc_placed = True

    # -- body ------------------------------------------------------------

    def _render_block(block: Block, *, top_level: bool = True) -> None:
        nonlocal body_started, toc_placed
        if isinstance(block, HeadingBlock):
            _add_heading(block)
        elif isinstance(block, ParagraphBlock):
            _body_paragraph(block.text, alignment=block.alignment)
            stats["paragraphs"] += 1
        elif isinstance(block, ListBlock):
            _add_list(block)
        elif isinstance(block, TableBlock):
            _add_table(block)
        elif isinstance(block, ImageBlock):
            _add_image(block)
        elif isinstance(block, ChartBlock):
            _add_chart(block)
        elif isinstance(block, CodeBlock):
            _add_code(block)
        elif isinstance(block, QuoteBlock):
            _add_quote(block)
        elif isinstance(block, CalloutBlock):
            _add_callout(block)
        elif isinstance(block, MetricsBlock):
            _add_metrics(block)
        elif isinstance(block, ChecklistBlock):
            _add_checklist(block)
        elif isinstance(block, MathBlock):
            _add_math(block)
        elif isinstance(block, DefinitionListBlock):
            _add_definition_list(block)
        elif isinstance(block, FootnotesBlock):
            _add_footnotes(block)
        elif isinstance(block, ColumnsBlock):
            # Word flows columns sequentially (graceful degradation).
            for col in block.columns:
                for nested in col:
                    if isinstance(nested, (PageBreakBlock, TocBlock)):
                        continue
                    _render_block(nested, top_level=False)
        elif isinstance(block, DividerBlock):
            _add_divider(doc, theme, qn=qn)
        elif isinstance(block, PageBreakBlock):
            doc.add_page_break()
            body_started = False
            return
        elif isinstance(block, SpacerBlock):
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(block.height_pt)
        elif isinstance(block, TocBlock):
            if top_level and not toc_placed:
                _add_toc(block.title)
                toc_placed = True
                body_started = False
            return
        if top_level:
            body_started = True

    for block in spec.blocks:
        _render_block(block)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    result = {
        "success": True,
        "output_path": str(output_path),
        "format": "docx",
        "file_size_bytes": output_path.stat().st_size,
        "content_stats": stats,
        "font_embedded": False,  # OOXML text relies on viewer fonts by design
        "east_asia_font": fonts.east_asia,
        "theme": theme.name,
        "toc": toc_placed,
        "warnings": warnings,
    }
    logger.info("docgen_docx_rendered", output_path=str(output_path), **stats)
    return result


# ---------------------------------------------------------------------------
# OOXML helpers
# ---------------------------------------------------------------------------


def _looks_chinese(spec: DocumentSpec) -> bool:
    sample = spec.title + "".join(
        getattr(b, "text", "") for b in spec.blocks[:8] if hasattr(b, "text")
    )
    return any("\u4e00" <= ch <= "\u9fff" for ch in sample)


def _apply_page_setup(section: Any, spec: DocumentSpec, wd_orient: Any, emu: Any) -> None:
    from docx.shared import Mm

    sizes_mm = {
        "A4": (210, 297),
        "LETTER": (216, 279),
        "LEGAL": (216, 356),
        "A3": (297, 420),
        "A5": (148, 210),
    }
    w_mm, h_mm = sizes_mm.get(spec.page.size, sizes_mm["A4"])
    if spec.page.orientation == "landscape":
        w_mm, h_mm = h_mm, w_mm
        section.orientation = wd_orient.LANDSCAPE
    section.page_width = Mm(w_mm)
    section.page_height = Mm(h_mm)
    margins = spec.page.margins
    pt_to_emu = 12700
    section.top_margin = emu(int(margins.top * pt_to_emu))
    section.bottom_margin = emu(int(margins.bottom * pt_to_emu))
    section.left_margin = emu(int(margins.left * pt_to_emu))
    section.right_margin = emu(int(margins.right * pt_to_emu))


def _apply_header_footer(
    section: Any,
    spec: DocumentSpec,
    theme: Theme,
    style_run: Any,
    align_map: dict[str, Any],
    *,
    qn: Any,
    pt: Any,
) -> None:
    for cfg, container in ((spec.header, section.header), (spec.footer, section.footer)):
        if cfg is None:
            continue
        p = container.paragraphs[0]
        p.alignment = align_map.get(cfg.alignment, align_map["center"])
        if cfg.text:
            run = p.add_run(cfg.text)
            style_run(
                run,
                font_name=theme.fonts.body,
                size=theme.sizes.small,
                color=theme.colors.text_light,
            )
        if cfg.show_page_number:
            if cfg.text:
                sep = p.add_run("  ·  ")
                style_run(
                    sep,
                    font_name=theme.fonts.body,
                    size=theme.sizes.small,
                    color=theme.colors.text_light,
                )
            run = p.add_run()
            style_run(
                run,
                font_name=theme.fonts.body,
                size=theme.sizes.small,
                color=theme.colors.text_light,
            )
            fld_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
            instr = run._element.makeelement(qn("w:instrText"), {})
            instr.text = "PAGE"
            fld_end = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
            run._element.append(fld_begin)
            run._element.append(instr)
            run._element.append(fld_end)


def _insert_toc_field(doc: Any, *, qn: Any) -> None:
    """Insert a native TOC field (levels 1-3, hyperlinked)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = run._element.makeelement(
        qn("w:fldChar"), {qn("w:fldCharType"): "begin"}
    )
    instr = run._element.makeelement(qn("w:instrText"), {})
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
    fld_end = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(fld_end)


def _shade_cell(cell: Any, hex_fill: str, *, qn: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_fill},
    )
    tc_pr.append(shd)


def _shade_run(run: Any, hex_fill: str, *, qn: Any) -> None:
    rpr = run._element.get_or_add_rPr()
    shd = rpr.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): _hex(hex_fill)},
    )
    rpr.append(shd)


def _set_cell_edge(
    cell: Any, edge: str, hex_color: str, *, qn: Any, size_eighth_pt: int
) -> None:
    """Set one border edge (``top``/``bottom``/…) on a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        tc_pr.append(borders)
    el = borders.find(qn(f"w:{edge}"))
    if el is not None:
        borders.remove(el)
    borders.append(
        borders.makeelement(
            qn(f"w:{edge}"),
            {
                qn("w:val"): "single",
                qn("w:sz"): str(max(2, size_eighth_pt)),
                qn("w:space"): "0",
                qn("w:color"): _hex(hex_color),
            },
        )
    )


def _set_table_rules(
    table: Any, pt: ProcessedTable, ts: TableStyleSpec, *, qn: Any
) -> None:
    """Apply the shared rule contract (header/row/total/outer) as cell borders."""

    def _row_edge(r_idx: int, edge: str, color: str, width_pt: float) -> None:
        sz = round(width_pt * 8)
        for cell in table.rows[r_idx].cells:
            _set_cell_edge(cell, edge, color, qn=qn, size_eighth_pt=sz)

    n_rows = len(table.rows)
    body_start = 1 if pt.has_header else 0

    if pt.has_header:
        _row_edge(0, "bottom", ts.header_rule, ts.header_rule_width)
    if ts.row_rule:
        last_ruled = n_rows - 1
        if pt.total_row_index is not None:
            last_ruled = body_start + pt.total_row_index - 1
        for r in range(body_start, last_ruled):
            _row_edge(r, "bottom", ts.row_rule, ts.row_rule_width)
    if pt.total_row_index is not None:
        _row_edge(body_start + pt.total_row_index, "top", ts.total_rule, ts.total_rule_width)
    if ts.outer_rule:
        _row_edge(0, "top", ts.outer_rule, ts.outer_rule_width)
        _row_edge(n_rows - 1, "bottom", ts.outer_rule, ts.outer_rule_width)


def _set_cell_left_border(
    cell: Any, hex_color: str, *, qn: Any, size_eighth_pt: int = 32
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        tc_pr.append(borders)
    left = borders.makeelement(
        qn("w:left"),
        {
            qn("w:val"): "single",
            qn("w:sz"): str(size_eighth_pt),
            qn("w:space"): "0",
            qn("w:color"): hex_color,
        },
    )
    borders.append(left)


def _add_divider(doc: Any, theme: Theme, *, qn: Any) -> None:
    p = doc.add_paragraph()
    p_pr = p._element.get_or_add_pPr()
    borders = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = borders.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): _hex(theme.colors.border),
        },
    )
    borders.append(bottom)
    p_pr.append(borders)


def _add_hyperlink(paragraph: Any, text: str, url: str, *, qn: Any) -> None:
    """Add a real hyperlink run (relationship-backed, styled blue underline)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = paragraph._element.makeelement(
        qn("w:hyperlink"), {qn("r:id"): r_id}
    )
    new_run = paragraph._element.makeelement(qn("w:r"), {})
    rpr = new_run.makeelement(qn("w:rPr"), {})
    color = rpr.makeelement(qn("w:color"), {qn("w:val"): "2E75B6"})
    underline = rpr.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rpr.append(color)
    rpr.append(underline)
    new_run.append(rpr)
    t = new_run.makeelement(qn("w:t"), {})
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
