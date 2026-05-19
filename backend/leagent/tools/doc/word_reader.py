"""Word Reader Tool - Extract text and structure from Word documents.

Uses python-docx for .docx/.docm (OOXML) and pyantiword (with optional
LibreOffice fallback) for legacy binary .doc (Word 97–2003).
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

import structlog

from leagent.tools.base import SyncTool, ToolCategory, ToolContext

logger = structlog.get_logger(__name__)


def _read_legacy_doc_libreoffice(file_path: Path) -> str | None:
    """Convert legacy .doc to plain text via headless LibreOffice."""
    exe = shutil.which("soffice") or shutil.which("libreoffice")
    if not exe:
        return None
    try:
        with tempfile.TemporaryDirectory() as td:
            proc = subprocess.run(
                [
                    exe,
                    "--headless",
                    "--nologo",
                    "--norestore",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    td,
                    str(file_path.resolve()),
                ],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            if proc.returncode != 0:
                logger.debug(
                    "libreoffice_doc_convert_failed",
                    returncode=proc.returncode,
                    stderr=(proc.stderr or "")[:500],
                )
                return None
            out_txt = Path(td) / f"{file_path.stem}.txt"
            if not out_txt.is_file():
                return None
            return out_txt.read_text(encoding="utf-8", errors="replace").strip()
    except (OSError, subprocess.TimeoutExpired) as e:
        logger.debug("libreoffice_doc_convert_error", error=str(e))
        return None


def _read_legacy_doc_binary(file_path: Path) -> tuple[str, str]:
    """Extract plain text from OLE Word .doc using pyantiword, then LibreOffice.

    Returns:
        (text, source) where source is ``pyantiword`` or ``libreoffice``.
    """
    try:
        from pyantiword.antiword_wrapper import extract_text_with_antiword

        raw = extract_text_with_antiword(str(file_path))
        if raw is not None:
            text = str(raw).strip()
            if text:
                return text, "pyantiword"
    except ImportError as e:
        raise RuntimeError(
            "pyantiword is not installed. Install with: pip install pyantiword"
        ) from e
    except Exception as e:
        logger.debug("pyantiword_extract_failed", path=str(file_path), error=str(e))

    lo_text = _read_legacy_doc_libreoffice(file_path)
    if lo_text:
        return lo_text, "libreoffice"

    raise ValueError(
        f"Could not read legacy Word document: {file_path}. "
        "The file may be corrupt, not a binary .doc (e.g. RTF renamed), or "
        "extraction failed; install LibreOffice (soffice) on the server as a fallback."
    )


def _build_legacy_doc_result(
    full_text: str,
    preserve_structure: bool,
    extract_tables: bool,
    extract_headers: bool,
    source: str,
) -> dict[str, Any]:
    """Shape legacy .doc output like the OOXML path (tables/headers not available)."""
    content_items: list[dict[str, Any]] = []
    if preserve_structure and full_text:
        for block in full_text.split("\n\n"):
            t = block.strip()
            if t:
                content_items.append({"type": "paragraph", "text": t})

    metadata = {
        "title": "",
        "author": "",
        "subject": "",
        "keywords": "",
        "created": "",
        "modified": "",
        "last_modified_by": "",
        "legacy_format": "doc",
        "extraction_source": source,
    }

    paragraph_count = len([i for i in content_items if i.get("type") == "paragraph"])
    if not content_items and full_text.strip():
        paragraph_count = 1

    result: dict[str, Any] = {
        "text": full_text,
        "total_characters": len(full_text),
        "paragraph_count": paragraph_count,
        "heading_count": 0,
        "table_count": 0,
        "metadata": metadata,
    }
    if preserve_structure and content_items:
        result["content"] = content_items
    if extract_tables:
        result["note"] = "Legacy .doc: table structure is not extracted; plain text only."
    if extract_headers:
        result["headers_footers"] = {"headers": [], "footers": []}
    return result


class WordReaderTool(SyncTool):
    """Extract text and structure from Word documents (.docx, .docm, .doc).

    Features:
    - .docx/.docm: paragraph extraction, headings, tables, optional headers/footers (python-docx)
    - .doc: plain text via pyantiword, with optional LibreOffice conversion fallback
    """

    name = "word_reader"
    description = (
        "Extract text from Word documents: .docx and .docm (structure, headings, tables) "
        "via python-docx; legacy binary .doc (Word 97–2003) as plain text via pyantiword "
        "with optional LibreOffice fallback."
    )
    category = ToolCategory.DOC
    version = "1.1.0"
    timeout_sec = 120
    aliases = ["docx", "word", "read_word", "docx_reader"]
    search_hint = "Word doc docx docm document read extract text headings paragraphs tables"
    is_concurrency_safe = True
    is_read_only = True
    interrupt_behavior = "cancel"
    max_result_size_chars = 200_000
    path_params = ("file_path",)

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the Word document (.docx, .docm, or legacy .doc).",
                },
                "extract_tables": {
                    "type": "boolean",
                    "description": "Whether to extract tables (.docx/.docm only). Defaults to True.",
                    "default": True,
                },
                "extract_headers": {
                    "type": "boolean",
                    "description": "Whether to extract header/footer content (.docx/.docm only). Defaults to False.",
                    "default": False,
                },
                "preserve_structure": {
                    "type": "boolean",
                    "description": "Whether to preserve structure (.docx/.docm); for .doc, splits on blank lines into paragraphs. Defaults to True.",
                    "default": True,
                },
            },
            "required": ["file_path"],
            "additionalProperties": False,
        }

    def get_activity_description(self, params: dict[str, Any] | None = None) -> str | None:
        return "Reading Word document"

    def execute_sync(self, params: dict[str, Any], context: ToolContext) -> dict[str, Any]:
        """Extract text and structure from a Word document."""
        file_path = Path(params["file_path"])
        extract_tables = params.get("extract_tables", True)
        extract_headers = params.get("extract_headers", False)
        preserve_structure = params.get("preserve_structure", True)

        if not file_path.exists():
            raise FileNotFoundError(f"Word document not found: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix == ".doc":
            logger.info("Reading legacy Word .doc", file_path=str(file_path))
            full_text, source = _read_legacy_doc_binary(file_path)
            result = _build_legacy_doc_result(
                full_text,
                preserve_structure,
                extract_tables,
                extract_headers,
                source,
            )
            logger.info(
                "Legacy Word document extraction complete",
                file_path=str(file_path),
                source=source,
                characters=result["total_characters"],
            )
            return result

        if suffix not in (".docx", ".docm"):
            raise ValueError(
                f"Unsupported file format: {file_path.suffix}. "
                "Supported: .doc (legacy), .docx, .docm."
            )

        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError as e:
            raise RuntimeError(
                "python-docx is not installed. Install with: pip install python-docx"
            ) from e

        logger.info("Opening Word document", file_path=str(file_path))

        try:
            doc = Document(str(file_path))
        except PackageNotFoundError as e:
            raise ValueError(f"Invalid or corrupted Word document: {file_path}") from e
        except Exception as e:
            raise RuntimeError(f"Failed to open Word document: {e}") from e

        content_items: list[dict[str, Any]] = []
        full_text_parts: list[str] = []
        heading_count = 0
        table_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"
            is_heading = style_name.startswith("Heading")

            if preserve_structure:
                item: dict[str, Any] = {
                    "type": "heading" if is_heading else "paragraph",
                    "text": text,
                }

                if is_heading:
                    heading_count += 1
                    try:
                        level = int(style_name.replace("Heading ", ""))
                    except ValueError:
                        level = 1
                    item["level"] = level

                content_items.append(item)
            else:
                full_text_parts.append(text)

        if preserve_structure:
            full_text_parts = [item["text"] for item in content_items]

        tables_data: list[list[list[str]]] = []
        if extract_tables:
            for table in doc.tables:
                table_data: list[list[str]] = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                    full_text_parts.append(" | ".join(row_data))

                tables_data.append(table_data)
                table_count += 1

                if preserve_structure:
                    content_items.append({
                        "type": "table",
                        "rows": table_data,
                        "row_count": len(table_data),
                        "col_count": len(table_data[0]) if table_data else 0,
                    })

        headers_footers: dict[str, list[str]] = {"headers": [], "footers": []}
        if extract_headers:
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        text = para.text.strip()
                        if text:
                            headers_footers["headers"].append(text)

                if section.footer:
                    for para in section.footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            headers_footers["footers"].append(text)

        full_text = "\n\n".join(full_text_parts)

        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
        }

        result: dict[str, Any] = {
            "text": full_text,
            "total_characters": len(full_text),
            "paragraph_count": len([i for i in content_items if i.get("type") == "paragraph"]),
            "heading_count": heading_count,
            "table_count": table_count,
            "metadata": metadata,
        }

        if preserve_structure:
            result["content"] = content_items

        if extract_tables and tables_data:
            result["tables"] = tables_data

        if extract_headers:
            result["headers_footers"] = headers_footers

        logger.info(
            "Word document extraction complete",
            file_path=str(file_path),
            paragraphs=result["paragraph_count"],
            headings=heading_count,
            tables=table_count,
        )

        return result
