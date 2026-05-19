"""Programmatic sample-file generation fixtures for LeAgent tests.

Covers every extension in EXTENSION_HANDLER_MAP plus the full image/archive
sets defined in config/constants.py.  All files are built in a session-scoped
temporary directory so expensive operations run only once per test session.

Generated types
───────────────
Documents : .pdf  .doc  .docx  .xls  .xlsx  .csv  .tsv  .md  .markdown
            .html .htm
Config    : .json .yaml .yml  .toml .ini  .cfg
Text/Code : .txt  .log  .py  .js  .ts  .java .c  .cpp .h  .rs  .go  .rb
            .sh   .sql  .xml .css
Archives  : .zip  .tar  .tar.gz  .tgz  .bz2  .xz
Images    : .png  .jpg  .jpeg  .bmp  .tiff .gif  .webp
"""

from __future__ import annotations

import csv
import gzip
import io
import json
import lzma
import struct
import tarfile
import zipfile
from pathlib import Path
from typing import Any

import pytest


# ---------------------------------------------------------------------------
# Session-scoped base directory
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def sample_dir(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Session-scoped directory that holds every sample file."""
    return tmp_path_factory.mktemp("samples")


# ===========================================================================
# Documents
# ===========================================================================


@pytest.fixture(scope="session")
def sample_pdf(sample_dir: Path) -> Path:
    """Multi-page PDF (PyMuPDF if available, else minimal valid PDF bytes)."""
    path = sample_dir / "sample.pdf"
    try:
        import fitz

        doc = fitz.open()
        for i in range(1, 4):
            page = doc.new_page()
            page.insert_text(
                (72, 72),
                f"Page {i}\nLeAgent test content – page {i}.\n"
                f"Fiscal Year 2024 Report\nTotal amount: ¥{i * 12345:,}",
            )
        doc.set_metadata({
            "title": "Test PDF Document",
            "author": "LeAgent Tests",
            "subject": "Unit Testing",
        })
        doc.save(str(path))
        doc.close()
    except ImportError:
        path.write_bytes(
            b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
            b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >>\nendobj\n"
            b"xref\n0 4\n0000000000 65535 f \n0000000009 00000 n \n"
            b"0000000058 00000 n \n0000000115 00000 n \n"
            b"trailer\n<< /Size 4 /Root 1 0 R >>\nstartxref\n190\n%%EOF\n"
        )
    return path


@pytest.fixture(scope="session")
def sample_docx(sample_dir: Path) -> Path:
    """DOCX with headings, paragraphs, and a table."""
    path = sample_dir / "sample.docx"
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("LeAgent Test Document", level=0)
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This document is generated programmatically for testing.")
        doc.add_paragraph("Second paragraph with more content and details.")
        doc.add_heading("Data Section", level=2)
        table = doc.add_table(rows=4, cols=3)
        headers = table.rows[0].cells
        headers[0].text, headers[1].text, headers[2].text = "Name", "Qty", "Price"
        for i, row in enumerate(table.rows[1:], 1):
            row.cells[0].text = f"Item {i}"
            row.cells[1].text = str(i * 10)
            row.cells[2].text = f"¥{i * 99.0:.2f}"
        doc.save(str(path))
    except ImportError:
        path.write_bytes(b"PK\x03\x04")
    return path


@pytest.fixture(scope="session")
def sample_doc(sample_dir: Path) -> Path:
    """File with .doc extension and RTF bytes (not OLE Word).

    Do not use for ``word_reader`` / antiword integration tests; use a real
    binary .doc or mock ``pyantiword`` instead.
    """
    path = sample_dir / "sample.doc"
    # Minimal RTF that many Word-compatible parsers accept
    path.write_bytes(
        b"{\\rtf1\\ansi\\deff0"
        b"{\\fonttbl{\\f0 Times New Roman;}}"
        b"\\f0\\fs24 LeAgent legacy DOC test content. "
        b"This file contains sample text for unit testing.}"
    )
    return path


@pytest.fixture(scope="session")
def sample_xlsx(sample_dir: Path) -> Path:
    """Multi-sheet XLSX workbook."""
    path = sample_dir / "sample.xlsx"
    try:
        import openpyxl

        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Employees"
        ws1.append(["Name", "Department", "Salary", "Start Date"])
        ws1.append(["Alice Zhang", "Engineering", 95000, "2020-03-01"])
        ws1.append(["Bob Li", "Marketing", 72000, "2021-06-15"])
        ws1.append(["Carol Wu", "Finance", 88000, "2019-11-20"])

        ws2 = wb.create_sheet("Expenses")
        ws2.append(["Category", "Q1", "Q2", "Q3", "Q4"])
        ws2.append(["Travel", 5000, 6200, 4800, 7100])
        ws2.append(["Office", 1200, 1300, 1250, 1400])
        ws2.append(["IT", 8000, 8000, 8000, 8000])

        ws3 = wb.create_sheet("Summary")
        ws3["A1"] = "Report Date"
        ws3["B1"] = "2024-01-01"
        ws3["A2"] = "Total Employees"
        ws3["B2"] = 3

        wb.save(str(path))
    except ImportError:
        path.write_bytes(b"PK\x03\x04")
    return path


@pytest.fixture(scope="session")
def sample_xls(sample_dir: Path) -> Path:
    """Legacy .xls file (xlwt if available, else minimal magic bytes)."""
    path = sample_dir / "sample.xls"
    try:
        import xlwt

        wb = xlwt.Workbook()
        ws = wb.add_sheet("Sheet1")
        headers = ["Name", "Value", "Date"]
        for col, h in enumerate(headers):
            ws.write(0, col, h)
        ws.write(1, 0, "Row1")
        ws.write(1, 1, 100)
        ws.write(1, 2, "2024-01-01")
        wb.save(str(path))
    except ImportError:
        # XLS BIFF8 magic bytes
        path.write_bytes(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1" + b"\x00" * 504)
    return path


@pytest.fixture(scope="session")
def sample_csv(sample_dir: Path) -> Path:
    """CSV with headers and mixed data types."""
    path = sample_dir / "sample.csv"
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["id", "name", "amount", "date", "approved"])
        writer.writerows([
            [1, "Expense Report A", 1200.50, "2024-01-15", "true"],
            [2, "Travel Reimbursement", 350.00, "2024-01-20", "false"],
            [3, "Office Supplies", 89.99, "2024-02-01", "true"],
            [4, "Conference Registration", 2500.00, "2024-02-10", "true"],
            [5, "Client Dinner", 485.75, "2024-02-15", "false"],
        ])
    return path


@pytest.fixture(scope="session")
def sample_tsv(sample_dir: Path) -> Path:
    """TSV (tab-separated) variant."""
    path = sample_dir / "sample.tsv"
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f, delimiter="\t")
        writer.writerow(["id", "product", "quantity", "price"])
        writer.writerows([
            [1, "Widget Pro", 100, 9.99],
            [2, "Gadget Max", 50, 24.99],
            [3, "Gizmo Lite", 200, 4.99],
        ])
    return path


@pytest.fixture(scope="session")
def sample_html(sample_dir: Path) -> Path:
    """HTML5 page with title, headings, table, links, and meta tags."""
    path = sample_dir / "sample.html"
    path.write_text(
        """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="description" content="LeAgent test HTML document">
  <meta name="keywords" content="test, leagent, automation">
  <title>LeAgent Test Page</title>
</head>
<body>
  <h1>Annual Report 2024</h1>
  <p>This document contains the <strong>quarterly results</strong> for FY 2024.
     See the <a href="https://example.com/details">full details</a>.</p>

  <h2>Financial Summary</h2>
  <table>
    <thead>
      <tr><th>Quarter</th><th>Revenue</th><th>Expenses</th><th>Net</th></tr>
    </thead>
    <tbody>
      <tr><td>Q1</td><td>¥1,200,000</td><td>¥850,000</td><td>¥350,000</td></tr>
      <tr><td>Q2</td><td>¥1,450,000</td><td>¥920,000</td><td>¥530,000</td></tr>
      <tr><td>Q3</td><td>¥1,380,000</td><td>¥880,000</td><td>¥500,000</td></tr>
      <tr><td>Q4</td><td>¥1,700,000</td><td>¥950,000</td><td>¥750,000</td></tr>
    </tbody>
  </table>

  <h2>Contacts</h2>
  <ul>
    <li><a href="mailto:cfo@example.com">CFO Office</a></li>
    <li><a href="https://example.com/ir">Investor Relations</a></li>
  </ul>
</body>
</html>""",
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_htm(sample_dir: Path) -> Path:
    """.htm alias – identical structure to .html."""
    path = sample_dir / "sample.htm"
    path.write_text(
        "<!DOCTYPE html><html><head><title>HTM Sample</title></head>"
        "<body><h1>HTM Test</h1><p>Simple HTM file for testing.</p></body></html>",
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_md(sample_dir: Path) -> Path:
    """Markdown with headings, code blocks, tables, and links."""
    path = sample_dir / "sample.md"
    path.write_text(
        """# LeAgent Test Document

## Introduction

This is a **Markdown** test file with _various_ formatting elements.

## Code Example

```python
def calculate_total(items: list[dict]) -> float:
    return sum(item["price"] * item["qty"] for item in items)
```

## Data Table

| Name    | Score | Grade |
|---------|-------|-------|
| Alice   | 95    | A     |
| Bob     | 82    | B     |
| Charlie | 71    | C     |

## Links

- [LeAgent Docs](https://docs.leagent.example.com)
- [GitHub](https://github.com/leagent)

## Nested List

1. First item
   - Sub-item A
   - Sub-item B
2. Second item
3. Third item
""",
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_markdown(sample_dir: Path) -> Path:
    """.markdown alias – same content as .md."""
    path = sample_dir / "sample.markdown"
    path.write_text(
        "# Markdown Extended Sample\n\n"
        "This file uses the `.markdown` extension.\n\n"
        "## Section\n\nContent for section.\n",
        encoding="utf-8",
    )
    return path


# ===========================================================================
# Config / structured text
# ===========================================================================


@pytest.fixture(scope="session")
def sample_json(sample_dir: Path) -> Path:
    """Nested JSON config file."""
    path = sample_dir / "config.json"
    data: dict[str, Any] = {
        "app": {"name": "leagent", "version": "1.0.0", "debug": False},
        "database": {"host": "localhost", "port": 5432, "name": "leagent_db"},
        "features": ["feature_a", "feature_b", "feature_c"],
        "limits": {"max_upload_mb": 100, "max_sessions": 500},
    }
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


@pytest.fixture(scope="session")
def sample_yaml(sample_dir: Path) -> Path:
    """Nested YAML config."""
    path = sample_dir / "config.yaml"
    path.write_text(
        """app:
  name: leagent
  version: "1.0.0"
  debug: false

database:
  host: localhost
  port: 5432
  name: leagent_db

redis:
  url: "redis://localhost:6379/0"
  max_connections: 50

features:
  - feature_a
  - feature_b
  - feature_c
""",
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_yml(sample_dir: Path) -> Path:
    """.yml alias."""
    path = sample_dir / "config.yml"
    path.write_text(
        "version: '3.9'\nservices:\n  app:\n    image: leagent:latest\n    ports:\n      - '7860:7860'\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_toml(sample_dir: Path) -> Path:
    """Nested TOML config."""
    path = sample_dir / "config.toml"
    path.write_text(
        """[app]
name = "leagent"
version = "1.0.0"
debug = false

[database]
host = "localhost"
port = 5432
name = "leagent_db"

[llm]
tier1_model = "Qwen2.5-32B-Instruct"
tier2_model = "Qwen2.5-7B-Instruct"

[features]
enabled = ["feature_a", "feature_b"]
""",
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_ini(sample_dir: Path) -> Path:
    """INI configuration file."""
    path = sample_dir / "config.ini"
    path.write_text(
        "[DEFAULT]\nDebug = False\nLogLevel = INFO\n\n"
        "[Database]\nHost = localhost\nPort = 5432\nName = leagent\n\n"
        "[Auth]\nJwtAlgorithm = HS256\nTokenExpireMinutes = 1440\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_cfg(sample_dir: Path) -> Path:
    """Generic .cfg configuration file."""
    path = sample_dir / "app.cfg"
    path.write_text(
        "# LeAgent config\n"
        "[server]\nhost = 0.0.0.0\nport = 7860\nworkers = 4\n\n"
        "[logging]\nlevel = INFO\nformat = json\n",
        encoding="utf-8",
    )
    return path


# ===========================================================================
# Plain text and code files
# ===========================================================================


@pytest.fixture(scope="session")
def sample_txt(sample_dir: Path) -> Path:
    """Multi-line plain text file."""
    path = sample_dir / "sample.txt"
    lines = [f"Line {i:03d}: the quick brown fox jumps over the lazy dog." for i in range(1, 31)]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


@pytest.fixture(scope="session")
def sample_log(sample_dir: Path) -> Path:
    """Structured application log file."""
    path = sample_dir / "app.log"
    path.write_text(
        '2024-01-15 09:00:01 INFO  [leagent.main] Application starting\n'
        '2024-01-15 09:00:02 INFO  [leagent.db] Database connected\n'
        '2024-01-15 09:00:03 INFO  [leagent.api] API listening on 0.0.0.0:7860\n'
        '2024-01-15 09:05:11 WARNING [leagent.auth] Failed login attempt user=unknown ip=192.168.1.100\n'
        '2024-01-15 09:10:22 ERROR [leagent.tools] Tool pdf_reader failed: file not found path=/tmp/x.pdf\n'
        '2024-01-15 09:15:30 INFO  [leagent.agent] Task completed task_id=abc-123 duration_ms=2450\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_py(sample_dir: Path) -> Path:
    """Python source file."""
    path = sample_dir / "example.py"
    path.write_text(
        '"""Example Python module for LeAgent testing."""\n\n'
        'from __future__ import annotations\n\n'
        'from typing import Any\n\n\n'
        'class ExpenseProcessor:\n'
        '    """Processes expense reports."""\n\n'
        '    def __init__(self, limit: float = 1000.0) -> None:\n'
        '        self.limit = limit\n\n'
        '    def approve(self, amount: float) -> bool:\n'
        '        """Return True if amount is within limit."""\n'
        '        return amount <= self.limit\n\n\n'
        'def calculate_total(items: list[dict[str, Any]]) -> float:\n'
        '    return sum(i["price"] * i.get("qty", 1) for i in items)\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_js(sample_dir: Path) -> Path:
    """JavaScript source file."""
    path = sample_dir / "example.js"
    path.write_text(
        '// LeAgent frontend helper\n\n'
        'const API_BASE = "http://localhost:7860/api/v1";\n\n'
        'async function runTask(query) {\n'
        '  const res = await fetch(`${API_BASE}/chat/run`, {\n'
        '    method: "POST",\n'
        '    headers: { "Content-Type": "application/json" },\n'
        '    body: JSON.stringify({ query }),\n'
        '  });\n'
        '  return res.json();\n'
        '}\n\n'
        'export { runTask };\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_ts(sample_dir: Path) -> Path:
    """TypeScript source file."""
    path = sample_dir / "example.ts"
    path.write_text(
        'interface TaskResult {\n'
        '  taskId: string;\n'
        '  status: "pending" | "running" | "completed" | "failed";\n'
        '  output?: string;\n'
        '}\n\n'
        'async function fetchResult(taskId: string): Promise<TaskResult> {\n'
        '  const response = await fetch(`/api/v1/tasks/${taskId}`);\n'
        '  return response.json() as TaskResult;\n'
        '}\n\n'
        'export { fetchResult, TaskResult };\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_java(sample_dir: Path) -> Path:
    """Java source file."""
    path = sample_dir / "Example.java"
    path.write_text(
        'package com.example.leagent;\n\n'
        '/**\n * Example Java class for LeAgent testing.\n */\n'
        'public class Example {\n\n'
        '    private final String name;\n\n'
        '    public Example(String name) {\n'
        '        this.name = name;\n'
        '    }\n\n'
        '    public String getName() {\n'
        '        return name;\n'
        '    }\n\n'
        '    public static void main(String[] args) {\n'
        '        System.out.println("LeAgent Java Test");\n'
        '    }\n'
        '}\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_c(sample_dir: Path) -> Path:
    """C source file."""
    path = sample_dir / "example.c"
    path.write_text(
        '#include <stdio.h>\n#include <stdlib.h>\n\n'
        '/* LeAgent C test file */\n\n'
        'typedef struct {\n'
        '    int id;\n'
        '    double amount;\n'
        '    char *description;\n'
        '} Expense;\n\n'
        'int main(void) {\n'
        '    Expense e = {1, 100.50, "Office supplies"};\n'
        '    printf("Expense %d: %.2f - %s\\n", e.id, e.amount, e.description);\n'
        '    return 0;\n'
        '}\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_cpp(sample_dir: Path) -> Path:
    """C++ source file."""
    path = sample_dir / "example.cpp"
    path.write_text(
        '#include <iostream>\n#include <vector>\n#include <string>\n\n'
        'class ExpenseManager {\n'
        'public:\n'
        '    void addExpense(const std::string& desc, double amount) {\n'
        '        expenses_.push_back({desc, amount});\n'
        '    }\n\n'
        '    double total() const {\n'
        '        double sum = 0;\n'
        '        for (const auto& e : expenses_) sum += e.second;\n'
        '        return sum;\n'
        '    }\n\n'
        'private:\n'
        '    std::vector<std::pair<std::string, double>> expenses_;\n'
        '};\n\n'
        'int main() {\n'
        '    std::cout << "LeAgent C++ Test" << std::endl;\n'
        '    return 0;\n'
        '}\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_h(sample_dir: Path) -> Path:
    """C header file."""
    path = sample_dir / "example.h"
    path.write_text(
        '#ifndef LEAGENT_EXAMPLE_H\n'
        '#define LEAGENT_EXAMPLE_H\n\n'
        '#include <stddef.h>\n\n'
        '/* Maximum expense amount */\n'
        '#define MAX_EXPENSE 100000.0\n\n'
        'typedef struct Expense {\n'
        '    int id;\n'
        '    double amount;\n'
        '    const char *category;\n'
        '} Expense;\n\n'
        'double calculate_total(const Expense *expenses, size_t count);\n\n'
        '#endif /* LEAGENT_EXAMPLE_H */\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_rs(sample_dir: Path) -> Path:
    """Rust source file."""
    path = sample_dir / "example.rs"
    path.write_text(
        '//! LeAgent Rust test module\n\n'
        'use std::collections::HashMap;\n\n'
        '#[derive(Debug, Clone)]\n'
        'pub struct Expense {\n'
        '    pub id: u32,\n'
        '    pub amount: f64,\n'
        '    pub category: String,\n'
        '}\n\n'
        'impl Expense {\n'
        '    pub fn new(id: u32, amount: f64, category: &str) -> Self {\n'
        '        Self { id, amount, category: category.to_owned() }\n'
        '    }\n\n'
        '    pub fn is_over_limit(&self, limit: f64) -> bool {\n'
        '        self.amount > limit\n'
        '    }\n'
        '}\n\n'
        'fn main() {\n'
        '    println!("LeAgent Rust Test");\n'
        '}\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_go(sample_dir: Path) -> Path:
    """Go source file."""
    path = sample_dir / "example.go"
    path.write_text(
        'package main\n\n'
        'import (\n'
        '\t"fmt"\n'
        '\t"math"\n'
        ')\n\n'
        '// Expense represents a single expense record\n'
        'type Expense struct {\n'
        '\tID       int\n'
        '\tAmount   float64\n'
        '\tCategory string\n'
        '}\n\n'
        '// IsOverBudget returns true if amount exceeds budget\n'
        'func (e Expense) IsOverBudget(budget float64) bool {\n'
        '\treturn e.Amount > budget\n'
        '}\n\n'
        'func main() {\n'
        '\te := Expense{1, math.Pi * 100, "travel"}\n'
        '\tfmt.Printf("LeAgent Go Test: %+v\\n", e)\n'
        '}\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_rb(sample_dir: Path) -> Path:
    """Ruby source file."""
    path = sample_dir / "example.rb"
    path.write_text(
        '# LeAgent Ruby test\n\n'
        'class Expense\n'
        '  attr_reader :id, :amount, :category\n\n'
        '  def initialize(id, amount, category)\n'
        '    @id = id\n'
        '    @amount = amount\n'
        '    @category = category\n'
        '  end\n\n'
        '  def over_limit?(limit)\n'
        '    @amount > limit\n'
        '  end\n'
        'end\n\n'
        'expense = Expense.new(1, 1200.50, "travel")\n'
        'puts "LeAgent Ruby Test: #{expense.inspect}"\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_sh(sample_dir: Path) -> Path:
    """Bash shell script."""
    path = sample_dir / "deploy.sh"
    path.write_text(
        '#!/usr/bin/env bash\n'
        '# LeAgent deployment script\n\n'
        'set -euo pipefail\n\n'
        'APP_NAME="leagent"\n'
        'DEPLOY_DIR="/opt/leagent"\n\n'
        'echo "Deploying ${APP_NAME}..."\n\n'
        'if [ ! -d "$DEPLOY_DIR" ]; then\n'
        '    mkdir -p "$DEPLOY_DIR"\n'
        'fi\n\n'
        'echo "Installing dependencies..."\n'
        'pip install -r requirements.txt\n\n'
        'echo "Starting services..."\n'
        'systemctl restart leagent\n\n'
        'echo "Deployment complete."\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_sql(sample_dir: Path) -> Path:
    """SQL DDL + DML script."""
    path = sample_dir / "schema.sql"
    path.write_text(
        '-- LeAgent database schema\n\n'
        'CREATE TABLE IF NOT EXISTS expenses (\n'
        '    id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
        '    submitter   VARCHAR(100) NOT NULL,\n'
        '    amount      NUMERIC(12, 2) NOT NULL,\n'
        '    category    VARCHAR(50) NOT NULL,\n'
        '    description TEXT,\n'
        '    status      VARCHAR(20) DEFAULT \'pending\',\n'
        '    created_at  TIMESTAMP WITH TIME ZONE DEFAULT NOW(),\n'
        '    updated_at  TIMESTAMP WITH TIME ZONE DEFAULT NOW()\n'
        ');\n\n'
        'CREATE INDEX idx_expenses_status ON expenses (status);\n'
        'CREATE INDEX idx_expenses_created ON expenses (created_at);\n\n'
        "INSERT INTO expenses (submitter, amount, category, description, status)\n"
        "VALUES\n"
        "    ('Alice Zhang', 1200.00, 'travel',  'Business trip to Shanghai', 'approved'),\n"
        "    ('Bob Li',      350.50,  'meals',   'Client dinner',              'pending'),\n"
        "    ('Carol Wu',    89.99,   'supplies', 'Printer cartridges',         'approved');\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_xml(sample_dir: Path) -> Path:
    """XML document with elements and attributes."""
    path = sample_dir / "data.xml"
    path.write_text(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<expenses version="1.0" generated="2024-01-15">\n'
        '  <expense id="1" status="approved">\n'
        '    <submitter>Alice Zhang</submitter>\n'
        '    <amount currency="CNY">1200.00</amount>\n'
        '    <category>travel</category>\n'
        '    <description>Business trip to Shanghai</description>\n'
        '  </expense>\n'
        '  <expense id="2" status="pending">\n'
        '    <submitter>Bob Li</submitter>\n'
        '    <amount currency="CNY">350.50</amount>\n'
        '    <category>meals</category>\n'
        '    <description>Client dinner</description>\n'
        '  </expense>\n'
        '</expenses>\n',
        encoding="utf-8",
    )
    return path


@pytest.fixture(scope="session")
def sample_css(sample_dir: Path) -> Path:
    """CSS stylesheet."""
    path = sample_dir / "styles.css"
    path.write_text(
        '/* LeAgent UI styles */\n\n'
        ':root {\n'
        '  --color-primary: #2563eb;\n'
        '  --color-bg: #f8fafc;\n'
        '  --font-sans: "Inter", sans-serif;\n'
        '}\n\n'
        'body {\n'
        '  font-family: var(--font-sans);\n'
        '  background: var(--color-bg);\n'
        '  margin: 0;\n'
        '  padding: 0;\n'
        '}\n\n'
        '.container {\n'
        '  max-width: 1200px;\n'
        '  margin: 0 auto;\n'
        '  padding: 0 1rem;\n'
        '}\n\n'
        '.btn-primary {\n'
        '  background: var(--color-primary);\n'
        '  color: #fff;\n'
        '  border: none;\n'
        '  border-radius: 6px;\n'
        '  padding: 0.5rem 1.25rem;\n'
        '  cursor: pointer;\n'
        '}\n',
        encoding="utf-8",
    )
    return path


# ===========================================================================
# Archives
# ===========================================================================


@pytest.fixture(scope="session")
def sample_zip(sample_dir: Path, sample_txt: Path, sample_csv: Path, sample_json: Path) -> Path:
    """ZIP archive containing text, CSV, and JSON files."""
    path = sample_dir / "sample.zip"
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(sample_txt, "docs/sample.txt")
        zf.write(sample_csv, "data/sample.csv")
        zf.write(sample_json, "config/config.json")
    return path


@pytest.fixture(scope="session")
def sample_tar(sample_dir: Path, sample_txt: Path, sample_json: Path) -> Path:
    """Uncompressed .tar archive."""
    path = sample_dir / "sample.tar"
    with tarfile.open(path, "w") as tf:
        tf.add(sample_txt, arcname="sample.txt")
        tf.add(sample_json, arcname="config.json")
    return path


@pytest.fixture(scope="session")
def sample_tar_gz(sample_dir: Path, sample_txt: Path, sample_csv: Path) -> Path:
    """Gzip-compressed .tar.gz archive."""
    path = sample_dir / "sample.tar.gz"
    with tarfile.open(path, "w:gz") as tf:
        tf.add(sample_txt, arcname="sample.txt")
        tf.add(sample_csv, arcname="data.csv")
    return path


@pytest.fixture(scope="session")
def sample_tgz(sample_dir: Path, sample_txt: Path) -> Path:
    """Gzip-compressed .tgz archive (same format as .tar.gz)."""
    path = sample_dir / "sample.tgz"
    with tarfile.open(path, "w:gz") as tf:
        tf.add(sample_txt, arcname="readme.txt")
    return path


@pytest.fixture(scope="session")
def sample_gz(sample_dir: Path, sample_txt: Path) -> Path:
    """Gzip-compressed single file (.gz, not a tar)."""
    path = sample_dir / "sample.txt.gz"
    with gzip.open(str(path), "wb") as gz:
        gz.write(sample_txt.read_bytes())
    return path


@pytest.fixture(scope="session")
def sample_bz2(sample_dir: Path, sample_txt: Path) -> Path:
    """Bzip2-compressed .tar.bz2 archive."""
    path = sample_dir / "sample.tar.bz2"
    with tarfile.open(path, "w:bz2") as tf:
        tf.add(sample_txt, arcname="sample.txt")
    return path


@pytest.fixture(scope="session")
def sample_xz(sample_dir: Path, sample_txt: Path) -> Path:
    """XZ-compressed .tar.xz archive."""
    path = sample_dir / "sample.tar.xz"
    with tarfile.open(path, "w:xz") as tf:
        tf.add(sample_txt, arcname="sample.txt")
    return path


# ===========================================================================
# Images
# ===========================================================================


def _pillow_image(path: Path, fmt: str, size: tuple[int, int] = (64, 64)) -> Path:
    """Write a small test image using Pillow (preferred) or raw bytes."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new("RGB", size, color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        draw.rectangle([4, 4, size[0] - 4, size[1] - 4], outline=(0, 120, 212), width=2)
        draw.text((8, 8), f"WA\n{fmt}", fill=(50, 50, 50))
        img.save(str(path), format=fmt)
    except ImportError:
        _write_raw_image(path, fmt)
    return path


def _write_raw_image(path: Path, fmt: str) -> None:
    """Fallback: write minimal valid image bytes for common formats."""
    if fmt.upper() == "PNG":
        path.write_bytes(_minimal_png())
    elif fmt.upper() in ("JPEG", "JPG"):
        path.write_bytes(_minimal_jpeg())
    elif fmt.upper() == "BMP":
        path.write_bytes(_minimal_bmp())
    elif fmt.upper() == "GIF":
        path.write_bytes(_minimal_gif())
    else:
        path.write_bytes(_minimal_png())


def _minimal_png() -> bytes:
    """Construct a 1×1 white PNG entirely in Python."""
    import zlib
    signature = b"\x89PNG\r\n\x1a\n"

    def _chunk(name: bytes, data: bytes) -> bytes:
        crc = struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + name + data + crc

    ihdr = _chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
    idat_data = zlib.compress(b"\x00\xff\xff\xff")
    idat = _chunk(b"IDAT", idat_data)
    iend = _chunk(b"IEND", b"")
    return signature + ihdr + idat + iend


def _minimal_jpeg() -> bytes:
    """Minimal 1×1 white JPEG."""
    return (
        b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00"
        b"\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t"
        b"\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a"
        b"\x1f\x1e\x1d\x1a\x1c\x1c $.' \",#\x1c\x1c(7),01444\x1f'9=82<.342\x1e\x1f"
        b"=49=86<78\xff\xc0\x00\x0b\x08\x00\x01\x00\x01\x01\x01\x11\x00"
        b"\xff\xc4\x00\x1f\x00\x00\x01\x05\x01\x01\x01\x01\x01\x01\x00"
        b"\x00\x00\x00\x00\x00\x00\x00\x01\x02\x03\x04\x05\x06\x07\x08"
        b"\t\n\x0b\xff\xc4\x00\xb5\x10\x00\x02\x01\x03\x03\x02\x04\x03"
        b"\x05\x05\x04\x04\x00\x00\x01}\x01\x02\x03\x00\x04\x11\x05\x12"
        b"!1A\x06\x13Qa\x07\"q\x142\x81\x91\xa1\x08#B\xb1\xc1\x15R\xd1"
        b"\xf0$3br\x82\t\n\x16\x17\x18\x19\x1a%&'()*456789:CDEFGHIJSTUVWXYZ"
        b"cdefghijstuvwxyz\x83\x84\x85\x86\x87\x88\x89\x8a\x92\x93\x94"
        b"\x95\x96\x97\x98\x99\x9a\xa2\xa3\xa4\xa5\xa6\xa7\xa8\xa9\xaa"
        b"\xb2\xb3\xb4\xb5\xb6\xb7\xb8\xb9\xba\xc2\xc3\xc4\xc5\xc6\xc7"
        b"\xc8\xc9\xca\xd2\xd3\xd4\xd5\xd6\xd7\xd8\xd9\xda\xe1\xe2\xe3"
        b"\xe4\xe5\xe6\xe7\xe8\xe9\xea\xf1\xf2\xf3\xf4\xf5\xf6\xf7\xf8"
        b"\xf9\xfa\xff\xda\x00\x08\x01\x01\x00\x00?\x00\xfb\xf4P\x00\x00"
        b"\x00\xff\xd9"
    )


def _minimal_bmp() -> bytes:
    """Minimal 2×2 24-bit BMP."""
    # Each row is padded to a multiple of 4 bytes; 2px × 3 bytes = 6 → pad to 8
    pixel_row = b"\xff\xff\xff\xff\xff\xff\x00\x00"  # 2 white pixels + 2 pad bytes
    pixel_data = pixel_row * 2  # 2 rows
    file_size = 54 + len(pixel_data)
    header = struct.pack(
        "<2sIHHI",
        b"BM", file_size, 0, 0, 54,
    )
    dib = struct.pack(
        "<IiiHHIIiiII",
        40, 2, 2, 1, 24, 0, len(pixel_data), 2835, 2835, 0, 0,
    )
    return header + dib + pixel_data


def _minimal_gif() -> bytes:
    """Minimal 1×1 GIF87a."""
    return (
        b"GIF87a"
        b"\x01\x00\x01\x00\x80\x00\x00"  # logical screen descriptor
        b"\xff\xff\xff\x00\x00\x00"       # global color table: white, black
        b"\x2c\x00\x00\x00\x00\x01\x00\x01\x00\x00"  # image descriptor
        b"\x02\x02\x4c\x01\x00"           # image data
        b"\x3b"                            # GIF trailer
    )


@pytest.fixture(scope="session")
def sample_png(sample_dir: Path) -> Path:
    return _pillow_image(sample_dir / "sample.png", "PNG")


@pytest.fixture(scope="session")
def sample_jpg(sample_dir: Path) -> Path:
    return _pillow_image(sample_dir / "sample.jpg", "JPEG")


@pytest.fixture(scope="session")
def sample_jpeg(sample_dir: Path) -> Path:
    """Alias to .jpeg extension."""
    path = sample_dir / "sample.jpeg"
    _pillow_image(path, "JPEG")
    return path


@pytest.fixture(scope="session")
def sample_bmp(sample_dir: Path) -> Path:
    return _pillow_image(sample_dir / "sample.bmp", "BMP")


@pytest.fixture(scope="session")
def sample_tiff(sample_dir: Path) -> Path:
    return _pillow_image(sample_dir / "sample.tiff", "TIFF")


@pytest.fixture(scope="session")
def sample_gif(sample_dir: Path) -> Path:
    try:
        from PIL import Image
        img = Image.new("P", (64, 64), color=0)
        img.save(str(sample_dir / "sample.gif"), format="GIF")
    except ImportError:
        (sample_dir / "sample.gif").write_bytes(_minimal_gif())
    return sample_dir / "sample.gif"


@pytest.fixture(scope="session")
def sample_webp(sample_dir: Path) -> Path:
    path = sample_dir / "sample.webp"
    try:
        from PIL import Image
        img = Image.new("RGB", (64, 64), color=(100, 149, 237))
        img.save(str(path), format="WEBP")
    except (ImportError, Exception):
        # Minimal RIFF/WEBP header with an empty VP8L chunk
        path.write_bytes(
            b"RIFF\x1c\x00\x00\x00WEBPVP8L\x10\x00\x00\x00"
            b"\x2f\x00\x00\x00\x00\x00\xfe\x03\x00\xfe\x03\x00\x00"
        )
    return path


# ===========================================================================
# Convenience: all-samples mapping
# ===========================================================================


@pytest.fixture(scope="session")
def all_samples(
    sample_dir: Path,
    sample_pdf: Path,
    sample_docx: Path,
    sample_doc: Path,
    sample_xlsx: Path,
    sample_xls: Path,
    sample_csv: Path,
    sample_tsv: Path,
    sample_html: Path,
    sample_htm: Path,
    sample_md: Path,
    sample_markdown: Path,
    sample_json: Path,
    sample_yaml: Path,
    sample_yml: Path,
    sample_toml: Path,
    sample_ini: Path,
    sample_cfg: Path,
    sample_txt: Path,
    sample_log: Path,
    sample_py: Path,
    sample_js: Path,
    sample_ts: Path,
    sample_java: Path,
    sample_c: Path,
    sample_cpp: Path,
    sample_h: Path,
    sample_rs: Path,
    sample_go: Path,
    sample_rb: Path,
    sample_sh: Path,
    sample_sql: Path,
    sample_xml: Path,
    sample_css: Path,
    sample_zip: Path,
    sample_tar: Path,
    sample_tar_gz: Path,
    sample_tgz: Path,
    sample_gz: Path,
    sample_bz2: Path,
    sample_xz: Path,
    sample_png: Path,
    sample_jpg: Path,
    sample_jpeg: Path,
    sample_bmp: Path,
    sample_tiff: Path,
    sample_gif: Path,
    sample_webp: Path,
) -> dict[str, Path]:
    """Return a mapping of extension → sample path for every supported type."""
    return {
        # Documents
        ".pdf": sample_pdf,
        ".docx": sample_docx,
        ".doc": sample_doc,
        ".xlsx": sample_xlsx,
        ".xls": sample_xls,
        ".csv": sample_csv,
        ".tsv": sample_tsv,
        ".html": sample_html,
        ".htm": sample_htm,
        ".md": sample_md,
        ".markdown": sample_markdown,
        # Config
        ".json": sample_json,
        ".yaml": sample_yaml,
        ".yml": sample_yml,
        ".toml": sample_toml,
        ".ini": sample_ini,
        ".cfg": sample_cfg,
        # Text / code
        ".txt": sample_txt,
        ".log": sample_log,
        ".py": sample_py,
        ".js": sample_js,
        ".ts": sample_ts,
        ".java": sample_java,
        ".c": sample_c,
        ".cpp": sample_cpp,
        ".h": sample_h,
        ".rs": sample_rs,
        ".go": sample_go,
        ".rb": sample_rb,
        ".sh": sample_sh,
        ".sql": sample_sql,
        ".xml": sample_xml,
        ".css": sample_css,
        # Archives
        ".zip": sample_zip,
        ".tar": sample_tar,
        ".tar.gz": sample_tar_gz,
        ".tgz": sample_tgz,
        ".gz": sample_gz,
        ".bz2": sample_bz2,
        ".xz": sample_xz,
        # Images
        ".png": sample_png,
        ".jpg": sample_jpg,
        ".jpeg": sample_jpeg,
        ".bmp": sample_bmp,
        ".tiff": sample_tiff,
        ".gif": sample_gif,
        ".webp": sample_webp,
    }
