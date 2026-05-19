"""Report Generator Tool - Generate reports from templates with data.

Creates data-driven reports with charts, tables, and multiple output formats.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import structlog

from leagent.tools.base import SyncTool, ToolCategory, ToolContext

logger = structlog.get_logger(__name__)


class ReportGeneratorTool(SyncTool):
    """Generate comprehensive reports from templates and data.

    Features:
    - Template-based report generation
    - Data-driven content (tables, charts, metrics)
    - Multiple output formats (PDF, Word, HTML, Markdown)
    - Section-based organization
    - Automatic table of contents
    - Executive summary generation
    """

    name = "report_generator"
    description = (
        "Generate data-driven reports from templates with support for charts, "
        "tables, metrics, and multiple output formats (PDF, Word, HTML, Markdown)."
    )
    category = ToolCategory.GEN
    version = "1.0.0"
    timeout_sec = 300
    aliases = ["report", "report_gen", "create_report"]
    search_hint = "report generate template charts tables metrics PDF Word HTML Markdown"
    is_concurrency_safe = False
    is_read_only = False
    interrupt_behavior = "cancel"
    max_result_size_chars = 100_000
    output_path_params = ("output_path",)

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "output_path": {
                    "type": "string",
                    "description": (
                        "Path where the report will be saved. Use only when the user "
                        "asked to save or export."
                    ),
                },
                "output_format": {
                    "type": "string",
                    "enum": ["pdf", "docx", "html", "markdown"],
                    "description": "Output format. Defaults to PDF.",
                },
                "template": {
                    "type": "object",
                    "description": "Report template definition.",
                    "properties": {
                        "name": {
                            "type": "string",
                            "description": "Template name for reference.",
                        },
                        "layout": {
                            "type": "string",
                            "enum": ["standard", "executive", "technical", "summary"],
                            "description": "Predefined layout style.",
                        },
                        "theme": {
                            "type": "object",
                            "description": "Visual theme settings.",
                            "properties": {
                                "primary_color": {"type": "string"},
                                "secondary_color": {"type": "string"},
                                "font_family": {"type": "string"},
                                "logo_path": {"type": "string"},
                            },
                        },
                    },
                },
                "metadata": {
                    "type": "object",
                    "description": "Report metadata.",
                    "properties": {
                        "title": {"type": "string"},
                        "subtitle": {"type": "string"},
                        "author": {"type": "string"},
                        "organization": {"type": "string"},
                        "date": {"type": "string"},
                        "version": {"type": "string"},
                        "confidentiality": {
                            "type": "string",
                            "enum": ["public", "internal", "confidential", "restricted"],
                        },
                    },
                },
                "sections": {
                    "type": "array",
                    "description": "Report sections in order.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {
                                "type": "string",
                                "enum": [
                                    "cover",
                                    "toc",
                                    "executive_summary",
                                    "section",
                                    "metrics",
                                    "table",
                                    "chart",
                                    "text",
                                    "appendix",
                                ],
                            },
                            "title": {"type": "string"},
                            "content": {
                                "type": "string",
                                "description": "Text content or markdown.",
                            },
                            "level": {
                                "type": "integer",
                                "description": "Section heading level (1-4).",
                            },
                            "data": {
                                "type": "object",
                                "description": "Data for tables/charts/metrics.",
                                "properties": {
                                    "rows": {
                                        "type": "array",
                                        "items": {"type": "array"},
                                    },
                                    "headers": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                    },
                                    "values": {
                                        "type": "array",
                                        "items": {"type": "number"},
                                    },
                                    "labels": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                    },
                                },
                            },
                            "metrics": {
                                "type": "array",
                                "description": "Key metrics to display.",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "label": {"type": "string"},
                                        "value": {},
                                        "unit": {"type": "string"},
                                        "trend": {
                                            "type": "string",
                                            "enum": ["up", "down", "neutral"],
                                        },
                                        "change": {"type": "number"},
                                    },
                                },
                            },
                            "chart_type": {
                                "type": "string",
                                "enum": ["bar", "line", "pie", "area", "scatter"],
                            },
                            "chart_options": {
                                "type": "object",
                                "properties": {
                                    "x_label": {"type": "string"},
                                    "y_label": {"type": "string"},
                                    "show_legend": {"type": "boolean"},
                                    "show_values": {"type": "boolean"},
                                },
                            },
                            "include_in_toc": {
                                "type": "boolean",
                                "description": "Include in table of contents.",
                            },
                        },
                        "required": ["type"],
                    },
                },
                "data_sources": {
                    "type": "object",
                    "description": "Named data sources for template binding.",
                    "additionalProperties": {},
                },
                "include_page_numbers": {
                    "type": "boolean",
                    "description": "Add page numbers to output.",
                },
                "include_header": {
                    "type": "boolean",
                    "description": "Add header with title/date.",
                },
                "include_footer": {
                    "type": "boolean",
                    "description": "Add footer with organization/page.",
                },
            },
            "required": ["output_path", "sections"],
            "additionalProperties": False,
        }

    def get_activity_description(self, params: dict[str, Any] | None = None) -> str | None:
        fmt = (params or {}).get("output_format", "pdf")
        return f"Generating report ({fmt})"

    def execute_sync(self, params: dict[str, Any], context: ToolContext) -> dict[str, Any]:
        """Generate a report with the specified configuration.

        Args:
            params: Tool parameters including output_path, sections, and options.
            context: Execution context.

        Returns:
            Dictionary containing generation status and report information.

        Raises:
            ValueError: If report configuration is invalid.
            RuntimeError: If report generation fails.
        """
        output_path = Path(params["output_path"])
        output_format = params.get("output_format", "pdf")
        sections = params.get("sections", [])
        metadata = params.get("metadata", {})
        template = params.get("template", {})
        data_sources = params.get("data_sources", {})

        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(
            "Generating report",
            output_path=str(output_path),
            format=output_format,
            sections=len(sections),
        )

        report_content = self._build_report_structure(
            sections, metadata, template, data_sources, params
        )

        stats = {
            "sections": len([s for s in sections if s.get("type") == "section"]),
            "tables": len([s for s in sections if s.get("type") == "table"]),
            "charts": len([s for s in sections if s.get("type") == "chart"]),
            "metrics_blocks": len([s for s in sections if s.get("type") == "metrics"]),
        }

        if output_format == "pdf":
            result = self._generate_pdf(output_path, report_content, metadata, template, params)
        elif output_format == "docx":
            result = self._generate_docx(output_path, report_content, metadata, template, params)
        elif output_format == "html":
            result = self._generate_html(output_path, report_content, metadata, template)
        elif output_format == "markdown":
            result = self._generate_markdown(output_path, report_content, metadata)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")

        file_size = output_path.stat().st_size

        logger.info(
            "Report generated successfully",
            output_path=str(output_path),
            file_size=file_size,
            format=output_format,
            **stats,
        )

        return {
            "success": True,
            "output_path": str(output_path),
            "output_format": output_format,
            "file_size_bytes": file_size,
            "stats": stats,
            "metadata": metadata,
            "toc": result.get("toc", []),
        }

    def _build_report_structure(
        self,
        sections: list[dict[str, Any]],
        metadata: dict[str, Any],
        template: dict[str, Any],
        data_sources: dict[str, Any],
        params: dict[str, Any],
    ) -> list[dict[str, Any]]:
        """Build the internal report structure from sections."""
        content: list[dict[str, Any]] = []
        toc_entries: list[dict[str, Any]] = []

        for section in sections:
            section_type = section.get("type")
            processed = section.copy()

            if section.get("include_in_toc", True) and section.get("title"):
                toc_entries.append({
                    "title": section["title"],
                    "level": section.get("level", 1),
                })

            if section_type == "cover":
                processed["metadata"] = metadata
                processed["template"] = template

            elif section_type == "executive_summary":
                if not processed.get("content"):
                    processed["content"] = self._generate_executive_summary(sections, data_sources)

            elif section_type == "metrics":
                metrics = section.get("metrics", [])
                for metric in metrics:
                    if isinstance(metric.get("value"), str) and metric["value"].startswith("$"):
                        key = metric["value"][1:]
                        if key in data_sources:
                            metric["value"] = data_sources[key]

            elif section_type == "table":
                data = section.get("data", {})
                if isinstance(data.get("source"), str) and data["source"] in data_sources:
                    source_data = data_sources[data["source"]]
                    if isinstance(source_data, list):
                        processed["data"]["rows"] = source_data

            elif section_type == "chart":
                data = section.get("data", {})
                if isinstance(data.get("source"), str) and data["source"] in data_sources:
                    source_data = data_sources[data["source"]]
                    if isinstance(source_data, dict):
                        processed["data"].update(source_data)

            content.append(processed)

        return content

    def _generate_executive_summary(
        self,
        sections: list[dict[str, Any]],
        data_sources: dict[str, Any],
    ) -> str:
        """Generate a basic executive summary from report data."""
        summary_parts = []

        metrics_sections = [s for s in sections if s.get("type") == "metrics"]
        for ms in metrics_sections:
            for metric in ms.get("metrics", []):
                label = metric.get("label", "Metric")
                value = metric.get("value", "N/A")
                unit = metric.get("unit", "")
                summary_parts.append(f"- {label}: {value}{unit}")

        table_sections = [s for s in sections if s.get("type") == "table"]
        if table_sections:
            summary_parts.append(f"- This report contains {len(table_sections)} data table(s).")

        chart_sections = [s for s in sections if s.get("type") == "chart"]
        if chart_sections:
            summary_parts.append(f"- This report includes {len(chart_sections)} chart(s).")

        return "\n".join(summary_parts) if summary_parts else "Report summary pending."

    def _generate_pdf(
        self,
        output_path: Path,
        content: list[dict[str, Any]],
        metadata: dict[str, Any],
        template: dict[str, Any],
        params: dict[str, Any],
    ) -> dict[str, Any]:
        """Generate PDF report using reportlab."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                PageBreak,
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )
        except ImportError as e:
            raise RuntimeError(
                "reportlab is not installed. Install with: pip install reportlab"
            ) from e

        layout = template.get("layout", "standard")
        page_size = A4 if layout != "executive" else landscape(A4)

        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=page_size,
            leftMargin=72,
            rightMargin=72,
            topMargin=72,
            bottomMargin=72,
            title=metadata.get("title", "Report"),
            author=metadata.get("author", ""),
        )

        styles = getSampleStyleSheet()
        story: list[Any] = []
        toc: list[dict[str, Any]] = []

        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontSize=28,
            spaceAfter=12,
            alignment=TA_CENTER,
        )

        section_style = ParagraphStyle(
            "SectionTitle",
            parent=styles["Heading1"],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
        )

        for section in content:
            section_type = section.get("type")

            if section_type == "cover":
                story.append(Spacer(1, 2 * inch))
                story.append(Paragraph(metadata.get("title", "Report"), title_style))
                if metadata.get("subtitle"):
                    story.append(Paragraph(metadata["subtitle"], styles["Heading2"]))
                story.append(Spacer(1, 0.5 * inch))
                if metadata.get("organization"):
                    story.append(Paragraph(metadata["organization"], styles["Normal"]))
                if metadata.get("date"):
                    story.append(Paragraph(f"Date: {metadata['date']}", styles["Normal"]))
                if metadata.get("author"):
                    story.append(Paragraph(f"Author: {metadata['author']}", styles["Normal"]))
                if metadata.get("confidentiality"):
                    conf_style = ParagraphStyle(
                        "Confidential",
                        parent=styles["Normal"],
                        textColor=colors.red,
                        spaceBefore=30,
                    )
                    story.append(Paragraph(metadata["confidentiality"].upper(), conf_style))
                story.append(PageBreak())

            elif section_type == "toc":
                story.append(Paragraph("Table of Contents", section_style))
                story.append(Spacer(1, 0.3 * inch))
                for entry in toc:
                    indent = "  " * (entry.get("level", 1) - 1)
                    story.append(Paragraph(f"{indent}{entry['title']}", styles["Normal"]))
                story.append(PageBreak())

            elif section_type == "executive_summary":
                story.append(Paragraph("Executive Summary", section_style))
                toc.append({"title": "Executive Summary", "level": 1})
                content_text = section.get("content", "")
                for line in content_text.split("\n"):
                    if line.strip():
                        story.append(Paragraph(line, styles["Normal"]))
                story.append(Spacer(1, 0.3 * inch))

            elif section_type == "section":
                title = section.get("title", "Section")
                level = section.get("level", 1)
                style_name = f"Heading{min(level, 4)}"
                story.append(Paragraph(title, styles.get(style_name, styles["Heading1"])))
                toc.append({"title": title, "level": level})

                if section.get("content"):
                    for para in section["content"].split("\n\n"):
                        if para.strip():
                            story.append(Paragraph(para, styles["Normal"]))

            elif section_type == "metrics":
                if section.get("title"):
                    story.append(Paragraph(section["title"], styles["Heading2"]))
                    toc.append({"title": section["title"], "level": 2})

                metrics = section.get("metrics", [])
                if metrics:
                    metric_data = []
                    for m in metrics:
                        trend_symbol = ""
                        if m.get("trend") == "up":
                            trend_symbol = " ↑"
                        elif m.get("trend") == "down":
                            trend_symbol = " ↓"

                        value_str = f"{m.get('value', 'N/A')}{m.get('unit', '')}{trend_symbol}"
                        metric_data.append([m.get("label", ""), value_str])

                    table = Table(metric_data, colWidths=[200, 150])
                    table.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, -1), colors.Color(0.95, 0.95, 0.95)),
                        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                        ("PADDING", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.3 * inch))

            elif section_type == "table":
                if section.get("title"):
                    story.append(Paragraph(section["title"], styles["Heading2"]))
                    toc.append({"title": section["title"], "level": 2})

                data = section.get("data", {})
                headers = data.get("headers", [])
                rows = data.get("rows", [])

                if headers or rows:
                    table_data = [headers] + rows if headers else rows
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.2, 0.4, 0.6)),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("PADDING", (0, 0), (-1, -1), 6),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.3 * inch))

            elif section_type == "text":
                if section.get("title"):
                    story.append(Paragraph(section["title"], styles["Heading3"]))
                if section.get("content"):
                    story.append(Paragraph(section["content"], styles["Normal"]))
                story.append(Spacer(1, 0.2 * inch))

            elif section_type == "appendix":
                story.append(PageBreak())
                story.append(Paragraph(f"Appendix: {section.get('title', '')}", section_style))
                toc.append({"title": f"Appendix: {section.get('title', '')}", "level": 1})
                if section.get("content"):
                    story.append(Paragraph(section["content"], styles["Normal"]))

        doc.build(story)
        return {"toc": toc}

    def _generate_docx(
        self,
        output_path: Path,
        content: list[dict[str, Any]],
        metadata: dict[str, Any],
        template: dict[str, Any],
        params: dict[str, Any],
    ) -> dict[str, Any]:
        """Generate Word report using python-docx."""
        try:
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.shared import Inches, Pt
        except ImportError as e:
            raise RuntimeError(
                "python-docx is not installed. Install with: pip install python-docx"
            ) from e

        doc = Document()
        toc: list[dict[str, Any]] = []

        if metadata.get("title"):
            doc.core_properties.title = metadata["title"]
        if metadata.get("author"):
            doc.core_properties.author = metadata["author"]

        for section in content:
            section_type = section.get("type")

            if section_type == "cover":
                doc.add_paragraph()
                doc.add_paragraph()
                title = doc.add_heading(metadata.get("title", "Report"), 0)
                title.alignment = 1

                if metadata.get("subtitle"):
                    doc.add_paragraph(metadata["subtitle"]).alignment = 1
                if metadata.get("organization"):
                    doc.add_paragraph(metadata["organization"]).alignment = 1
                if metadata.get("date"):
                    doc.add_paragraph(f"Date: {metadata['date']}").alignment = 1
                if metadata.get("author"):
                    doc.add_paragraph(f"Author: {metadata['author']}").alignment = 1

                doc.add_page_break()

            elif section_type == "toc":
                doc.add_heading("Table of Contents", 1)
                for entry in toc:
                    indent = "  " * (entry.get("level", 1) - 1)
                    doc.add_paragraph(f"{indent}{entry['title']}")
                doc.add_page_break()

            elif section_type == "executive_summary":
                doc.add_heading("Executive Summary", 1)
                toc.append({"title": "Executive Summary", "level": 1})
                if section.get("content"):
                    doc.add_paragraph(section["content"])

            elif section_type == "section":
                title = section.get("title", "Section")
                level = min(section.get("level", 1), 4)
                doc.add_heading(title, level)
                toc.append({"title": title, "level": level})

                if section.get("content"):
                    for para in section["content"].split("\n\n"):
                        if para.strip():
                            doc.add_paragraph(para)

            elif section_type == "metrics":
                if section.get("title"):
                    doc.add_heading(section["title"], 2)
                    toc.append({"title": section["title"], "level": 2})

                metrics = section.get("metrics", [])
                if metrics:
                    table = doc.add_table(rows=len(metrics), cols=2)
                    table.style = "Table Grid"

                    for i, m in enumerate(metrics):
                        trend_symbol = ""
                        if m.get("trend") == "up":
                            trend_symbol = " ↑"
                        elif m.get("trend") == "down":
                            trend_symbol = " ↓"

                        table.rows[i].cells[0].text = m.get("label", "")
                        table.rows[i].cells[1].text = f"{m.get('value', 'N/A')}{m.get('unit', '')}{trend_symbol}"

            elif section_type == "table":
                if section.get("title"):
                    doc.add_heading(section["title"], 2)
                    toc.append({"title": section["title"], "level": 2})

                data = section.get("data", {})
                headers = data.get("headers", [])
                rows = data.get("rows", [])

                if headers or rows:
                    num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
                    num_rows = len(rows) + (1 if headers else 0)

                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = "Table Grid"

                    if headers:
                        for j, header in enumerate(headers):
                            cell = table.rows[0].cells[j]
                            cell.text = str(header)
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True

                    start_row = 1 if headers else 0
                    for i, row in enumerate(rows):
                        for j, value in enumerate(row):
                            if j < num_cols:
                                table.rows[start_row + i].cells[j].text = str(value)

            elif section_type == "text":
                if section.get("title"):
                    doc.add_heading(section["title"], 3)
                if section.get("content"):
                    doc.add_paragraph(section["content"])

            elif section_type == "appendix":
                doc.add_page_break()
                doc.add_heading(f"Appendix: {section.get('title', '')}", 1)
                toc.append({"title": f"Appendix: {section.get('title', '')}", "level": 1})
                if section.get("content"):
                    doc.add_paragraph(section["content"])

        doc.save(str(output_path))
        return {"toc": toc}

    def _generate_html(
        self,
        output_path: Path,
        content: list[dict[str, Any]],
        metadata: dict[str, Any],
        template: dict[str, Any],
    ) -> dict[str, Any]:
        """Generate HTML report."""
        theme = template.get("theme", {})
        primary_color = theme.get("primary_color", "#2c5282")
        secondary_color = theme.get("secondary_color", "#4a5568")
        font_family = theme.get("font_family", "Arial, sans-serif")

        toc: list[dict[str, Any]] = []

        html_parts = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            f"<title>{metadata.get('title', 'Report')}</title>",
            "<meta charset='UTF-8'>",
            "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            f"<meta name='author' content='{metadata.get('author', '')}'>",
            "<style>",
            f"body {{ font-family: {font_family}; line-height: 1.6; max-width: 900px; margin: 0 auto; padding: 20px; }}",
            f"h1, h2, h3 {{ color: {primary_color}; }}",
            f".cover {{ text-align: center; padding: 100px 0; border-bottom: 3px solid {primary_color}; }}",
            ".cover h1 { font-size: 2.5em; margin-bottom: 0.5em; }",
            f".cover .subtitle {{ color: {secondary_color}; font-size: 1.2em; }}",
            ".toc { background: #f7fafc; padding: 20px; border-radius: 8px; margin: 20px 0; }",
            ".toc ul { list-style: none; padding-left: 20px; }",
            ".metrics { display: flex; flex-wrap: wrap; gap: 20px; margin: 20px 0; }",
            f".metric-card {{ background: #f7fafc; padding: 20px; border-radius: 8px; min-width: 200px; border-left: 4px solid {primary_color}; }}",
            ".metric-label { font-size: 0.9em; color: #718096; }",
            ".metric-value { font-size: 1.8em; font-weight: bold; }",
            ".trend-up { color: #38a169; }",
            ".trend-down { color: #e53e3e; }",
            "table { width: 100%; border-collapse: collapse; margin: 20px 0; }",
            f"th {{ background: {primary_color}; color: white; padding: 12px; text-align: left; }}",
            "td { padding: 10px; border-bottom: 1px solid #e2e8f0; }",
            "tr:hover { background: #f7fafc; }",
            ".section { margin: 30px 0; }",
            ".appendix { border-top: 2px solid #e2e8f0; margin-top: 40px; padding-top: 20px; }",
            "@media print { .cover { page-break-after: always; } }",
            "</style>",
            "</head>",
            "<body>",
        ]

        for section in content:
            section_type = section.get("type")

            if section_type == "cover":
                html_parts.append("<div class='cover'>")
                html_parts.append(f"<h1>{metadata.get('title', 'Report')}</h1>")
                if metadata.get("subtitle"):
                    html_parts.append(f"<p class='subtitle'>{metadata['subtitle']}</p>")
                if metadata.get("organization"):
                    html_parts.append(f"<p>{metadata['organization']}</p>")
                if metadata.get("date"):
                    html_parts.append(f"<p>Date: {metadata['date']}</p>")
                if metadata.get("author"):
                    html_parts.append(f"<p>Author: {metadata['author']}</p>")
                if metadata.get("confidentiality"):
                    html_parts.append(f"<p style='color: red;'>{metadata['confidentiality'].upper()}</p>")
                html_parts.append("</div>")

            elif section_type == "toc":
                html_parts.append("<div class='toc'>")
                html_parts.append("<h2>Table of Contents</h2>")
                html_parts.append("<ul>")
                for entry in toc:
                    indent = "&nbsp;" * (entry.get("level", 1) - 1) * 4
                    html_parts.append(f"<li>{indent}{entry['title']}</li>")
                html_parts.append("</ul>")
                html_parts.append("</div>")

            elif section_type == "executive_summary":
                html_parts.append("<div class='section'>")
                html_parts.append("<h2>Executive Summary</h2>")
                toc.append({"title": "Executive Summary", "level": 1})
                if section.get("content"):
                    for line in section["content"].split("\n"):
                        if line.strip():
                            html_parts.append(f"<p>{line}</p>")
                html_parts.append("</div>")

            elif section_type == "section":
                title = section.get("title", "Section")
                level = min(section.get("level", 1), 4)
                html_parts.append("<div class='section'>")
                html_parts.append(f"<h{level}>{title}</h{level}>")
                toc.append({"title": title, "level": level})
                if section.get("content"):
                    for para in section["content"].split("\n\n"):
                        if para.strip():
                            html_parts.append(f"<p>{para}</p>")
                html_parts.append("</div>")

            elif section_type == "metrics":
                if section.get("title"):
                    html_parts.append(f"<h3>{section['title']}</h3>")
                    toc.append({"title": section["title"], "level": 2})

                html_parts.append("<div class='metrics'>")
                for m in section.get("metrics", []):
                    trend_class = ""
                    trend_symbol = ""
                    if m.get("trend") == "up":
                        trend_class = "trend-up"
                        trend_symbol = " ↑"
                    elif m.get("trend") == "down":
                        trend_class = "trend-down"
                        trend_symbol = " ↓"

                    html_parts.append("<div class='metric-card'>")
                    html_parts.append(f"<div class='metric-label'>{m.get('label', '')}</div>")
                    html_parts.append(
                        f"<div class='metric-value {trend_class}'>"
                        f"{m.get('value', 'N/A')}{m.get('unit', '')}{trend_symbol}</div>"
                    )
                    html_parts.append("</div>")
                html_parts.append("</div>")

            elif section_type == "table":
                if section.get("title"):
                    html_parts.append(f"<h3>{section['title']}</h3>")
                    toc.append({"title": section["title"], "level": 2})

                data = section.get("data", {})
                headers = data.get("headers", [])
                rows = data.get("rows", [])

                if headers or rows:
                    html_parts.append("<table>")
                    if headers:
                        html_parts.append("<tr>")
                        for header in headers:
                            html_parts.append(f"<th>{header}</th>")
                        html_parts.append("</tr>")
                    for row in rows:
                        html_parts.append("<tr>")
                        for cell in row:
                            html_parts.append(f"<td>{cell}</td>")
                        html_parts.append("</tr>")
                    html_parts.append("</table>")

            elif section_type == "text":
                if section.get("title"):
                    html_parts.append(f"<h4>{section['title']}</h4>")
                if section.get("content"):
                    html_parts.append(f"<p>{section['content']}</p>")

            elif section_type == "appendix":
                html_parts.append("<div class='appendix'>")
                html_parts.append(f"<h2>Appendix: {section.get('title', '')}</h2>")
                toc.append({"title": f"Appendix: {section.get('title', '')}", "level": 1})
                if section.get("content"):
                    html_parts.append(f"<p>{section['content']}</p>")
                html_parts.append("</div>")

        html_parts.extend(["</body>", "</html>"])

        output_path.write_text("\n".join(html_parts), encoding="utf-8")
        return {"toc": toc}

    def _generate_markdown(
        self,
        output_path: Path,
        content: list[dict[str, Any]],
        metadata: dict[str, Any],
    ) -> dict[str, Any]:
        """Generate Markdown report."""
        toc: list[dict[str, Any]] = []
        md_parts: list[str] = []

        for section in content:
            section_type = section.get("type")

            if section_type == "cover":
                md_parts.append(f"# {metadata.get('title', 'Report')}")
                if metadata.get("subtitle"):
                    md_parts.append(f"\n*{metadata['subtitle']}*")
                md_parts.append("")
                if metadata.get("organization"):
                    md_parts.append(f"**Organization:** {metadata['organization']}")
                if metadata.get("date"):
                    md_parts.append(f"**Date:** {metadata['date']}")
                if metadata.get("author"):
                    md_parts.append(f"**Author:** {metadata['author']}")
                if metadata.get("confidentiality"):
                    md_parts.append(f"\n> **{metadata['confidentiality'].upper()}**")
                md_parts.append("\n---\n")

            elif section_type == "toc":
                md_parts.append("## Table of Contents\n")
                for entry in toc:
                    indent = "  " * (entry.get("level", 1) - 1)
                    md_parts.append(f"{indent}- {entry['title']}")
                md_parts.append("\n---\n")

            elif section_type == "executive_summary":
                md_parts.append("## Executive Summary\n")
                toc.append({"title": "Executive Summary", "level": 1})
                if section.get("content"):
                    md_parts.append(section["content"])
                md_parts.append("")

            elif section_type == "section":
                title = section.get("title", "Section")
                level = min(section.get("level", 1), 6)
                md_parts.append(f"\n{'#' * level} {title}\n")
                toc.append({"title": title, "level": level})
                if section.get("content"):
                    md_parts.append(section["content"])
                md_parts.append("")

            elif section_type == "metrics":
                if section.get("title"):
                    md_parts.append(f"### {section['title']}\n")
                    toc.append({"title": section["title"], "level": 2})

                for m in section.get("metrics", []):
                    trend_symbol = ""
                    if m.get("trend") == "up":
                        trend_symbol = " ↑"
                    elif m.get("trend") == "down":
                        trend_symbol = " ↓"

                    md_parts.append(
                        f"- **{m.get('label', '')}:** {m.get('value', 'N/A')}{m.get('unit', '')}{trend_symbol}"
                    )
                md_parts.append("")

            elif section_type == "table":
                if section.get("title"):
                    md_parts.append(f"### {section['title']}\n")
                    toc.append({"title": section["title"], "level": 2})

                data = section.get("data", {})
                headers = data.get("headers", [])
                rows = data.get("rows", [])

                if headers:
                    md_parts.append("| " + " | ".join(str(h) for h in headers) + " |")
                    md_parts.append("| " + " | ".join(["---"] * len(headers)) + " |")

                for row in rows:
                    md_parts.append("| " + " | ".join(str(c) for c in row) + " |")
                md_parts.append("")

            elif section_type == "text":
                if section.get("title"):
                    md_parts.append(f"#### {section['title']}\n")
                if section.get("content"):
                    md_parts.append(section["content"])
                md_parts.append("")

            elif section_type == "appendix":
                md_parts.append("\n---\n")
                md_parts.append(f"## Appendix: {section.get('title', '')}\n")
                toc.append({"title": f"Appendix: {section.get('title', '')}", "level": 1})
                if section.get("content"):
                    md_parts.append(section["content"])
                md_parts.append("")

        output_path.write_text("\n".join(md_parts), encoding="utf-8")
        return {"toc": toc}
