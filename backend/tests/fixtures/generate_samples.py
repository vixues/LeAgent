"""Standalone sample-file generator.

Generates every file type that LeAgent's tools and file-processing service
support into a persistent ``_cache/`` directory (or a custom path).  Running
this once creates all binary and text samples without network access.

Usage
─────
    python tests/fixtures/generate_samples.py            # generate to _cache/
    python tests/fixtures/generate_samples.py --out /tmp/wa_samples
    python tests/fixtures/generate_samples.py --list     # list generated files
"""

from __future__ import annotations

import argparse
import csv
import gzip
import json
import logging
import lzma
import struct
import sys
import tarfile
import zipfile
import zlib
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

DEFAULT_OUT = Path(__file__).parent / "_cache"


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _minimal_png_bytes() -> bytes:
    """Build a 4×4 white PNG entirely in Python (no external libraries)."""
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(name: bytes, data: bytes) -> bytes:
        crc = struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + name + data + crc

    # 4×4 8-bit RGB image
    raw = b""
    for _ in range(4):
        raw += b"\x00" + b"\xff\xff\xff" * 4  # filter byte + 4 white pixels
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 4, 4, 8, 2, 0, 0, 0))
    idat = chunk(b"IDAT", zlib.compress(raw))
    iend = chunk(b"IEND", b"")
    return sig + ihdr + idat + iend


def _minimal_jpeg_bytes() -> bytes:
    return (
        b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00"
        b"\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t"
        b"\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a"
        b"\x1f\x1e\x1d\x1a\x1c\x1c $.' \",#\x1c\x1c(7),01444\x1f'9=82<.342\x1e\x1f"
        b"=49=86<78\xff\xc0\x00\x0b\x08\x00\x01\x00\x01\x01\x01\x11\x00"
        b"\xff\xc4\x00\x1f\x00\x00\x01\x05\x01\x01\x01\x01\x01\x01\x00"
        b"\x00\x00\x00\x00\x00\x00\x00\x01\x02\x03\x04\x05\x06\x07\x08"
        b"\t\n\x0b\xff\xc4\x00\xb5\x10\x00\x02\x01\x03\x03\x02\x04\x03"
        b"\x05\x05\x04\x04\x00\x00\x01}\x01\x02\x03\x00\x04\x11\x05\x12"
        b"!1A\x06\x13Qa\x07\"q\x142\x81\x91\xa1\x08#B\xb1\xc1\x15R\xd1"
        b"\xf0$3br\x82\t\n\x16\x17\x18\x19\x1a%&'()*456789:CDEFGHIJSTUVWXYZ"
        b"cdefghijstuvwxyz\x83\x84\x85\x86\x87\x88\x89\x8a\x92\x93\x94"
        b"\x95\x96\x97\x98\x99\x9a\xa2\xa3\xa4\xa5\xa6\xa7\xa8\xa9\xaa"
        b"\xb2\xb3\xb4\xb5\xb6\xb7\xb8\xb9\xba\xc2\xc3\xc4\xc5\xc6\xc7"
        b"\xc8\xc9\xca\xd2\xd3\xd4\xd5\xd6\xd7\xd8\xd9\xda\xe1\xe2\xe3"
        b"\xe4\xe5\xe6\xe7\xe8\xe9\xea\xf1\xf2\xf3\xf4\xf5\xf6\xf7\xf8"
        b"\xf9\xfa\xff\xda\x00\x08\x01\x01\x00\x00?\x00\xfb\xf4P\x00\x00"
        b"\x00\xff\xd9"
    )


def _minimal_bmp_bytes() -> bytes:
    """4×4 24-bit BMP."""
    row = b"\xff\xff\xff" * 4  # 4 white pixels, 12 bytes
    pixel_data = row * 4  # 4 rows = 48 bytes (12 bytes per row, no padding needed since 12%4==0)
    file_size = 54 + len(pixel_data)
    header = struct.pack("<2sIHHI", b"BM", file_size, 0, 0, 54)
    dib = struct.pack("<IiiHHIIiiII", 40, 4, 4, 1, 24, 0, len(pixel_data), 2835, 2835, 0, 0)
    return header + dib + pixel_data


def _minimal_gif_bytes() -> bytes:
    return (
        b"GIF87a\x01\x00\x01\x00\x80\x00\x00"
        b"\xff\xff\xff\x00\x00\x00"
        b"\x2c\x00\x00\x00\x00\x01\x00\x01\x00\x00"
        b"\x02\x02\x4c\x01\x00\x3b"
    )


def _make_image(path: Path, fmt: str) -> None:
    """Write a small image using Pillow when available, otherwise raw bytes."""
    try:
        from PIL import Image, ImageDraw
        img = Image.new("RGB", (80, 80), color=(240, 248, 255))
        draw = ImageDraw.Draw(img)
        draw.rectangle([4, 4, 76, 76], outline=(0, 102, 204), width=3)
        draw.text((10, 30), f"WA\n.{fmt.lower()}", fill=(20, 20, 20))
        save_fmt = "JPEG" if fmt.upper() in ("JPG", "JPEG") else fmt.upper()
        img.save(str(path), format=save_fmt)
    except (ImportError, Exception) as exc:
        logger.debug("Pillow unavailable (%s); using raw bytes for %s", exc, path.name)
        raw = {
            "PNG": _minimal_png_bytes,
            "JPG": _minimal_jpeg_bytes,
            "JPEG": _minimal_jpeg_bytes,
            "BMP": _minimal_bmp_bytes,
            "GIF": _minimal_gif_bytes,
        }.get(fmt.upper(), _minimal_png_bytes)()
        path.write_bytes(raw)


# ---------------------------------------------------------------------------
# Generators
# ---------------------------------------------------------------------------


def generate_all(out_dir: Path) -> dict[str, Path]:
    """Generate every sample file into *out_dir* and return ext → path mapping."""
    out_dir.mkdir(parents=True, exist_ok=True)
    files: dict[str, Path] = {}

    # ── PDF ─────────────────────────────────────────────────────────────────
    pdf = out_dir / "sample.pdf"
    try:
        import fitz
        doc = fitz.open()
        for i in range(1, 4):
            page = doc.new_page()
            page.insert_text(
                (72, 72),
                f"Page {i} – LeAgent Test Document\n\n"
                f"Fiscal Year 2024 – Section {i}\n"
                "Total expenses: ¥1,200,000\n"
                "Approved by: Finance Director\n",
            )
        doc.set_metadata({"title": "Test PDF", "author": "LeAgent"})
        doc.save(str(pdf))
        doc.close()
    except ImportError:
        pdf.write_bytes(
            b"%PDF-1.4\n1 0 obj\n<</Type/Catalog/Pages 2 0 R>>\nendobj\n"
            b"2 0 obj\n<</Type/Pages/Kids[3 0 R]/Count 1>>\nendobj\n"
            b"3 0 obj\n<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>\nendobj\n"
            b"xref\n0 4\n0000000000 65535 f \n0000000009 00000 n \n"
            b"0000000058 00000 n \n0000000115 00000 n \n"
            b"trailer\n<</Size 4/Root 1 0 R>>\nstartxref\n190\n%%EOF\n"
        )
    files[".pdf"] = pdf
    logger.info("✓ .pdf")

    # ── DOCX ────────────────────────────────────────────────────────────────
    docx = out_dir / "sample.docx"
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("LeAgent Annual Report 2024", level=0)
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(
            "This report summarises the financial performance for fiscal year 2024. "
            "Total revenue increased by 23% compared to the prior year."
        )
        doc.add_heading("Expense Summary", level=2)
        table = doc.add_table(rows=4, cols=3)
        headers = ["Category", "Amount (¥)", "Status"]
        for col, h in enumerate(headers):
            table.rows[0].cells[col].text = h
        data = [("Travel", "120,000", "Approved"), ("IT Equipment", "85,000", "Pending"),
                ("Training", "32,000", "Approved")]
        for row, (cat, amt, st) in zip(table.rows[1:], data):
            row.cells[0].text = cat
            row.cells[1].text = amt
            row.cells[2].text = st
        doc.save(str(docx))
    except ImportError:
        docx.write_bytes(b"PK\x03\x04")
    files[".docx"] = docx
    logger.info("✓ .docx")

    # ── XLSX ────────────────────────────────────────────────────────────────
    xlsx = out_dir / "sample.xlsx"
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Employees"
        ws.append(["ID", "Name", "Department", "Salary", "Start Date"])
        rows = [
            (1, "Alice Zhang", "Engineering", 95000, "2020-03-01"),
            (2, "Bob Li",      "Marketing",   72000, "2021-06-15"),
            (3, "Carol Wu",    "Finance",      88000, "2019-11-20"),
            (4, "David Wang",  "Operations",  65000, "2022-02-10"),
            (5, "Eve Chen",    "HR",           70000, "2021-09-05"),
        ]
        for r in rows:
            ws.append(r)
        ws2 = wb.create_sheet("Expenses")
        ws2.append(["Category", "Q1", "Q2", "Q3", "Q4", "Total"])
        for cat, q1, q2, q3, q4 in [
            ("Travel",  5000, 6200, 4800, 7100),
            ("Office",  1200, 1300, 1250, 1400),
            ("IT",      8000, 8000, 8000, 8000),
        ]:
            ws2.append([cat, q1, q2, q3, q4, q1 + q2 + q3 + q4])
        wb.save(str(xlsx))
    except ImportError:
        xlsx.write_bytes(b"PK\x03\x04")
    files[".xlsx"] = xlsx
    logger.info("✓ .xlsx")

    # ── XLS ─────────────────────────────────────────────────────────────────
    xls = out_dir / "sample.xls"
    try:
        import xlwt
        wb = xlwt.Workbook()
        ws = wb.add_sheet("Data")
        for col, h in enumerate(["ID", "Name", "Value"]):
            ws.write(0, col, h)
        ws.write(1, 0, 1); ws.write(1, 1, "Alpha"); ws.write(1, 2, 100)
        ws.write(2, 0, 2); ws.write(2, 1, "Beta");  ws.write(2, 2, 200)
        wb.save(str(xls))
    except ImportError:
        xls.write_bytes(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1" + b"\x00" * 504)
    files[".xls"] = xls
    logger.info("✓ .xls")

    # ── CSV ─────────────────────────────────────────────────────────────────
    csv_path = out_dir / "sample.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["id", "name", "amount", "date", "approved", "category"])
        for i, (name, amt, date, appr, cat) in enumerate([
            ("Alice Zhang",     1200.50, "2024-01-15", "true",  "travel"),
            ("Bob Li",           350.00, "2024-01-20", "false", "meals"),
            ("Carol Wu",          89.99, "2024-02-01", "true",  "supplies"),
            ("David Wang",      2500.00, "2024-02-10", "true",  "conference"),
            ("Eve Chen",         485.75, "2024-02-15", "false", "meals"),
            ("Frank Zhang",     3200.00, "2024-03-01", "true",  "equipment"),
        ], 1):
            w.writerow([i, name, amt, date, appr, cat])
    files[".csv"] = csv_path
    logger.info("✓ .csv")

    # ── TSV ─────────────────────────────────────────────────────────────────
    tsv = out_dir / "sample.tsv"
    with tsv.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter="\t")
        w.writerow(["id", "product", "quantity", "unit_price", "total"])
        for i, (prod, qty, price) in enumerate([
            ("Widget Pro", 100, 9.99),
            ("Gadget Max",  50, 24.99),
            ("Gizmo Lite", 200, 4.99),
            ("Tool Kit",    30, 49.99),
        ], 1):
            w.writerow([i, prod, qty, price, qty * price])
    files[".tsv"] = tsv
    logger.info("✓ .tsv")

    # ── HTML ────────────────────────────────────────────────────────────────
    html = out_dir / "sample.html"
    html.write_text(
        '<!DOCTYPE html>\n<html lang="zh-CN">\n<head>\n'
        '  <meta charset="utf-8">\n'
        '  <title>LeAgent 2024年度报告</title>\n'
        '</head>\n<body>\n'
        '  <h1>财务年度报告</h1>\n'
        '  <p>本报告总结了 <strong>2024财年</strong> 的财务绩效。</p>\n'
        '  <h2>收支摘要</h2>\n'
        '  <table border="1">\n'
        '    <thead><tr><th>季度</th><th>收入</th><th>支出</th><th>净利润</th></tr></thead>\n'
        '    <tbody>\n'
        '      <tr><td>Q1</td><td>¥1,200,000</td><td>¥850,000</td><td>¥350,000</td></tr>\n'
        '      <tr><td>Q2</td><td>¥1,450,000</td><td>¥920,000</td><td>¥530,000</td></tr>\n'
        '      <tr><td>Q3</td><td>¥1,380,000</td><td>¥880,000</td><td>¥500,000</td></tr>\n'
        '      <tr><td>Q4</td><td>¥1,700,000</td><td>¥950,000</td><td>¥750,000</td></tr>\n'
        '    </tbody>\n'
        '  </table>\n'
        '  <p><a href="https://example.com">公司官网</a></p>\n'
        '</body>\n</html>\n',
        encoding="utf-8",
    )
    files[".html"] = html
    logger.info("✓ .html")

    # ── Markdown ────────────────────────────────────────────────────────────
    md = out_dir / "sample.md"
    md.write_text(
        "# LeAgent Test Document\n\n"
        "## Introduction\n\nThis is a **Markdown** test file.\n\n"
        "## Code Example\n\n```python\ndef greet(name: str) -> str:\n    return f'Hello, {name}'\n```\n\n"
        "## Data Table\n\n"
        "| Name  | Score | Grade |\n|-------|-------|-------|\n"
        "| Alice | 95    | A     |\n| Bob   | 82    | B     |\n\n"
        "## Lists\n\n1. First\n2. Second\n   - Sub A\n   - Sub B\n",
        encoding="utf-8",
    )
    files[".md"] = md
    logger.info("✓ .md")

    # ── Config files ─────────────────────────────────────────────────────────
    json_file = out_dir / "config.json"
    json_file.write_text(
        json.dumps({
            "app": {"name": "leagent", "version": "1.0.0", "debug": False},
            "database": {"host": "localhost", "port": 5432},
            "features": ["feature_a", "feature_b"],
        }, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    files[".json"] = json_file
    logger.info("✓ .json")

    yaml_file = out_dir / "config.yaml"
    yaml_file.write_text(
        "app:\n  name: leagent\n  version: '1.0.0'\n  debug: false\n\n"
        "database:\n  host: localhost\n  port: 5432\n\n"
        "features:\n  - feature_a\n  - feature_b\n",
        encoding="utf-8",
    )
    files[".yaml"] = yaml_file
    logger.info("✓ .yaml")

    toml_file = out_dir / "config.toml"
    toml_file.write_text(
        '[app]\nname = "leagent"\nversion = "1.0.0"\ndebug = false\n\n'
        '[database]\nhost = "localhost"\nport = 5432\n\n'
        '[features]\nenabled = ["feature_a", "feature_b"]\n',
        encoding="utf-8",
    )
    files[".toml"] = toml_file
    logger.info("✓ .toml")

    ini_file = out_dir / "config.ini"
    ini_file.write_text(
        "[DEFAULT]\nDebug = False\nLogLevel = INFO\n\n"
        "[Database]\nHost = localhost\nPort = 5432\n\n"
        "[Auth]\nJwtAlgorithm = HS256\n",
        encoding="utf-8",
    )
    files[".ini"] = ini_file
    logger.info("✓ .ini")

    # ── Text / code files ────────────────────────────────────────────────────
    _TEXT_FILES: dict[str, str] = {
        "sample.txt":   "\n".join(f"Line {i:03d}: The quick brown fox." for i in range(1, 31)),
        "app.log":      (
            "2024-01-15 09:00:01 INFO  Application starting\n"
            "2024-01-15 09:00:02 INFO  Database connected\n"
            "2024-01-15 09:05:11 WARNING Failed login attempt\n"
            "2024-01-15 09:10:22 ERROR  Tool execution failed\n"
        ),
        "example.py":   (
            '"""Python example."""\nfrom __future__ import annotations\n\n'
            'class Processor:\n    def run(self, data: dict) -> dict:\n        return data\n'
        ),
        "example.js":   (
            "// JS example\nconst API = 'http://localhost:7860/api/v1';\n\n"
            "async function run(q) {\n  const r = await fetch(`${API}/chat`);\n  return r.json();\n}\n"
        ),
        "example.ts":   (
            "interface Result { id: string; status: string; }\n\n"
            "async function fetch_result(id: string): Promise<Result> {\n"
            "  const r = await fetch(`/api/v1/tasks/${id}`);\n  return r.json();\n}\n"
        ),
        "Example.java":  (
            "package com.example;\npublic class Example {\n"
            "    public static void main(String[] args) {\n"
            "        System.out.println(\"LeAgent Java Test\");\n    }\n}\n"
        ),
        "example.c":    (
            "#include <stdio.h>\nint main(void) {\n    printf(\"LeAgent C Test\\n\");\n    return 0;\n}\n"
        ),
        "example.cpp":  (
            "#include <iostream>\nint main() {\n    std::cout << \"LeAgent C++ Test\" << std::endl;\n    return 0;\n}\n"
        ),
        "example.h":    (
            "#ifndef EXAMPLE_H\n#define EXAMPLE_H\n#define MAX_VALUE 100\ntypedef struct { int id; } Record;\n#endif\n"
        ),
        "example.rs":   (
            "fn main() {\n    println!(\"LeAgent Rust Test\");\n}\n"
        ),
        "example.go":   (
            'package main\nimport "fmt"\nfunc main() { fmt.Println("LeAgent Go Test") }\n'
        ),
        "example.rb":   (
            "class Greeter\n  def greet(name) = \"Hello, #{name}\"\nend\n"
            "puts Greeter.new.greet('LeAgent')\n"
        ),
        "deploy.sh":    (
            "#!/bin/bash\nset -euo pipefail\necho 'Deploying LeAgent...'\npip install -r requirements.txt\n"
            "echo 'Done.'\n"
        ),
        "schema.sql":   (
            "CREATE TABLE IF NOT EXISTS expenses (\n"
            "    id SERIAL PRIMARY KEY,\n    amount NUMERIC(12,2),\n"
            "    category VARCHAR(50),\n    created_at TIMESTAMP DEFAULT NOW()\n);\n"
            "INSERT INTO expenses (amount, category) VALUES (100.00, 'travel'), (50.00, 'meals');\n"
        ),
        "data.xml":     (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<records>\n  <record id="1"><name>Alice</name><value>100</value></record>\n'
            '  <record id="2"><name>Bob</name><value>200</value></record>\n</records>\n'
        ),
        "styles.css":   (
            ":root { --primary: #2563eb; }\n"
            "body { font-family: sans-serif; margin: 0; }\n"
            ".btn { background: var(--primary); color: #fff; padding: .5rem 1rem; }\n"
        ),
    }

    ext_map = {
        ".txt": "sample.txt", ".log": "app.log", ".py": "example.py",
        ".js": "example.js",  ".ts": "example.ts", ".java": "Example.java",
        ".c": "example.c",    ".cpp": "example.cpp", ".h": "example.h",
        ".rs": "example.rs",  ".go": "example.go",   ".rb": "example.rb",
        ".sh": "deploy.sh",   ".sql": "schema.sql",  ".xml": "data.xml",
        ".css": "styles.css",
    }

    for ext, fname in ext_map.items():
        p = out_dir / fname
        p.write_text(_TEXT_FILES[fname], encoding="utf-8")
        files[ext] = p
        logger.info("✓ %s", ext)

    # ── Archives ─────────────────────────────────────────────────────────────
    txt_path = files[".txt"]
    csv_path = files[".csv"]
    json_path = files[".json"]

    zip_path = out_dir / "sample.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(txt_path, "docs/sample.txt")
        zf.write(csv_path, "data/sample.csv")
        zf.write(json_path, "config/config.json")
    files[".zip"] = zip_path
    logger.info("✓ .zip")

    tar_path = out_dir / "sample.tar"
    with tarfile.open(tar_path, "w") as tf:
        tf.add(txt_path, arcname="sample.txt")
    files[".tar"] = tar_path
    logger.info("✓ .tar")

    tgz_path = out_dir / "sample.tar.gz"
    with tarfile.open(tgz_path, "w:gz") as tf:
        tf.add(txt_path, arcname="sample.txt")
        tf.add(csv_path, arcname="data.csv")
    files[".tar.gz"] = tgz_path
    files[".tgz"]    = tgz_path
    logger.info("✓ .tar.gz / .tgz")

    gz_path = out_dir / "sample.txt.gz"
    with gzip.open(str(gz_path), "wb") as gz:
        gz.write(txt_path.read_bytes())
    files[".gz"] = gz_path
    logger.info("✓ .gz")

    bz2_path = out_dir / "sample.tar.bz2"
    with tarfile.open(bz2_path, "w:bz2") as tf:
        tf.add(txt_path, arcname="sample.txt")
    files[".bz2"] = bz2_path
    logger.info("✓ .bz2")

    xz_path = out_dir / "sample.tar.xz"
    with tarfile.open(xz_path, "w:xz") as tf:
        tf.add(txt_path, arcname="sample.txt")
    files[".xz"] = xz_path
    logger.info("✓ .xz")

    # ── Images ───────────────────────────────────────────────────────────────
    for ext, fmt in [
        (".png",  "PNG"),
        (".jpg",  "JPEG"),
        (".jpeg", "JPEG"),
        (".bmp",  "BMP"),
        (".tiff", "TIFF"),
        (".gif",  "GIF"),
        (".webp", "WEBP"),
    ]:
        fname = f"sample{ext}"
        p = out_dir / fname
        _make_image(p, fmt)
        files[ext] = p
        logger.info("✓ %s", ext)

    return files


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)-8s %(message)s")

    parser = argparse.ArgumentParser(description="Generate LeAgent test sample files")
    parser.add_argument("--out", default=str(DEFAULT_OUT), help="Output directory")
    parser.add_argument("--list", action="store_true", help="List files that would be generated")
    args = parser.parse_args()

    out_dir = Path(args.out)

    if args.list:
        print("File types that will be generated:")
        for ext in sorted([
            ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".tsv",
            ".html", ".htm", ".md", ".markdown",
            ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg",
            ".txt", ".log", ".py", ".js", ".ts", ".java", ".c", ".cpp",
            ".h", ".rs", ".go", ".rb", ".sh", ".sql", ".xml", ".css",
            ".zip", ".tar", ".tar.gz", ".tgz", ".gz", ".bz2", ".xz",
            ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif", ".webp",
        ]):
            print(f"  {ext}")
        sys.exit(0)

    print(f"Generating sample files to: {out_dir}")
    files = generate_all(out_dir)

    print(f"\n{'─' * 60}")
    print(f"Generated {len(files)} sample files:")
    print(f"{'─' * 60}")
    for ext, path in sorted(files.items()):
        size = path.stat().st_size
        print(f"  {ext:12s}  {path.name:30s}  {size:>8,} bytes")
    print(f"{'─' * 60}")
    print(f"Output directory: {out_dir.resolve()}")
