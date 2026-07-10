"""Checklist capability: model, stats, fence parsing, per-format rendering, tool."""

from __future__ import annotations

import json
from typing import TYPE_CHECKING

import pytest

from leagent.docgen.checklist import (
    build_checklist_block,
    checklist_stats,
    checklist_to_dict,
    parse_workflow_checklist,
)
from leagent.docgen.markdown import parse_markdown_blocks
from leagent.docgen.model import ChecklistBlock, DocumentSpec

if TYPE_CHECKING:
    from pathlib import Path


def _block() -> ChecklistBlock:
    return ChecklistBlock.model_validate(
        {
            "title": "Launch",
            "description": "Readiness for v2.0",
            "groups": [
                {
                    "name": "Engineering",
                    "items": [
                        {"text": "API", "status": "completed", "priority": "high"},
                        {
                            "text": "Migrate",
                            "status": "in_progress",
                            "assignee": "amy",
                            "due_date": "2026-08-01",
                            "sub_items": [
                                {"text": "write", "status": "completed"},
                                {"text": "review", "status": "pending"},
                            ],
                        },
                        {"text": "Load test", "status": "blocked", "priority": "critical"},
                    ],
                },
                {"name": "Docs", "items": [{"text": "Guide", "status": "skipped"}]},
            ],
        }
    )


# ---------------------------------------------------------------------------
# Model + stats
# ---------------------------------------------------------------------------


def test_status_coercion() -> None:
    b = ChecklistBlock.model_validate(
        {"items": [{"text": "a", "status": "In Progress"}, {"text": "b", "status": "BLOCKED"}]}
    )
    assert [i.status for i in b.items] == ["in_progress", "blocked"]


def test_flat_items_normalized_to_group() -> None:
    b = ChecklistBlock.model_validate({"items": [{"text": "a"}]})
    groups = b.normalized_groups()
    assert len(groups) == 1
    assert groups[0].items[0].text == "a"
    assert groups[0].items[0].status == "pending"


def test_stats_counts_subitems_and_excludes_skipped_from_progress() -> None:
    stats = checklist_stats(_block())
    # 3 + 2 sub + 1 = 6 total
    assert stats["total_items"] == 6
    assert stats["completed"] == 2
    assert stats["in_progress"] == 1
    assert stats["blocked"] == 1
    assert stats["skipped"] == 1
    # 2 completed / (6 - 1 skipped) = 40%
    assert stats["progress_percentage"] == 40


def test_checklist_to_dict_has_stats_and_groups() -> None:
    data = checklist_to_dict(_block())
    assert "stats" in data
    assert isinstance(data["groups"], list) and len(data["groups"]) == 2
    assert "items" not in data  # flat items folded into groups


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------


def test_checklist_fence() -> None:
    md = (
        "```checklist\n"
        '{"title": "T", "items": [{"text": "a", "status": "completed"}, '
        '{"text": "b", "priority": "high"}]}\n'
        "```\n"
    )
    (block,) = parse_markdown_blocks(md)
    assert isinstance(block, ChecklistBlock)
    assert block.title == "T"
    assert block.items[0].status == "completed"
    assert block.items[1].priority == "high"


def test_invalid_checklist_fence_degrades_to_code() -> None:
    from leagent.docgen.model import CodeBlock

    (block,) = parse_markdown_blocks("```checklist\nnot json {{{\n```\n")
    assert isinstance(block, CodeBlock)


def test_workflow_source_adapter(tmp_path: Path) -> None:
    wf = tmp_path / "wf.yaml"
    wf.write_text(
        "name: Deploy\nnodes:\n  - {id: a, label: Build}\n  - {id: b, name: Test}\n",
        encoding="utf-8",
    )
    block = parse_workflow_checklist(wf)
    assert block.title == "Deploy"
    assert [i.text for i in block.items] == ["Build", "Test"]
    assert all(i.status == "pending" for i in block.items)


def test_build_checklist_block_manual_with_overrides() -> None:
    block = build_checklist_block(
        {
            "title": "T",
            "items": [{"text": "x"}],
            "include_progress": False,
            "include_legend": False,
        }
    )
    assert block.title == "T"
    assert block.show_progress is False
    assert block.show_legend is False


# ---------------------------------------------------------------------------
# Rendering (all formats)
# ---------------------------------------------------------------------------


def _spec() -> DocumentSpec:
    return DocumentSpec.model_validate({"title": "CL", "blocks": [_block()]})


def test_render_pdf_checklist(tmp_path: Path) -> None:
    from leagent.docgen.renderers.pdf import render_pdf

    result = render_pdf(_spec(), tmp_path / "cl.pdf")
    assert result["success"] is True
    fitz = pytest.importorskip("fitz")
    doc = fitz.open(str(tmp_path / "cl.pdf"))
    try:
        text = doc[0].get_text()
        assert "Launch" in text
        assert "Engineering" in text
        assert "Migrate" in text
        assert "@amy" in text
        assert "40%" in text  # progress
    finally:
        doc.close()


def test_render_docx_checklist(tmp_path: Path) -> None:
    import docx as docx_lib

    from leagent.docgen.renderers.docx import render_docx

    result = render_docx(_spec(), tmp_path / "cl.docx")
    assert result["success"] is True
    d = docx_lib.Document(str(tmp_path / "cl.docx"))
    all_text = "\n".join(p.text for p in d.paragraphs)
    assert "Launch" in all_text and "Engineering" in all_text
    assert "@amy" in all_text
    assert len(d.tables) >= 1  # progress bar table


def test_render_html_checklist(tmp_path: Path) -> None:
    from leagent.docgen.renderers.html import render_html

    result = render_html(_spec(), tmp_path / "cl.html")
    assert result["success"] is True
    text = (tmp_path / "cl.html").read_text(encoding="utf-8")
    assert 'class="checklist"' in text
    assert "cl-progress-bar" in text
    assert "cl-prio" in text
    assert "cl-legend" in text


def test_render_markdown_checklist_roundtrips(tmp_path: Path) -> None:
    from leagent.docgen.renderers.html import render_markdown

    result = render_markdown(_spec(), tmp_path / "cl.md")
    assert result["success"] is True
    text = (tmp_path / "cl.md").read_text(encoding="utf-8")
    assert "- [x] API" in text
    assert "- [~] Migrate" in text  # in_progress marker
    assert "  - [x] write" in text  # nested sub-item
    # Reparses to a task list (GFM).
    from leagent.docgen.model import ListBlock

    assert any(isinstance(b, ListBlock) for b in parse_markdown_blocks(text))


# ---------------------------------------------------------------------------
# Tool
# ---------------------------------------------------------------------------


class _Ctx:
    session_id = "test-session"


def test_tool_json_export(tmp_path: Path) -> None:
    from leagent.tools.gen.checklist_tool import ChecklistGeneratorTool

    tool = ChecklistGeneratorTool()
    out = tmp_path / "c.json"
    result = tool.execute_sync(
        {
            "output_path": str(out),
            "title": "Q",
            "items": [
                {"text": "a", "status": "completed"},
                {"text": "b", "status": "pending"},
            ],
        },
        _Ctx(),
    )
    assert result["success"] is True
    data = json.loads(out.read_text(encoding="utf-8"))
    assert data["stats"]["total_items"] == 2
    assert data["stats"]["progress_percentage"] == 50


def test_tool_aliases_include_legacy_name() -> None:
    from leagent.tools.gen.checklist_tool import ChecklistGeneratorTool

    assert ChecklistGeneratorTool.name == "checklist_generate"
    assert "checklist_generator" in ChecklistGeneratorTool.aliases


def test_tool_pdf_and_docx(tmp_path: Path) -> None:
    from leagent.tools.gen.checklist_tool import ChecklistGeneratorTool

    tool = ChecklistGeneratorTool()
    params = {
        "title": "R",
        "groups": [{"name": "G", "items": [{"text": "x", "status": "completed"}]}],
    }
    r_pdf = tool.execute_sync({**params, "output_path": str(tmp_path / "c.pdf")}, _Ctx())
    assert r_pdf["success"] is True
    r_docx = tool.execute_sync({**params, "output_path": str(tmp_path / "c.docx")}, _Ctx())
    assert r_docx["success"] is True


def test_tool_empty_checklist_raises(tmp_path: Path) -> None:
    from leagent.tools.gen.checklist_tool import ChecklistGeneratorTool

    tool = ChecklistGeneratorTool()
    with pytest.raises(ValueError, match="empty"):
        tool.execute_sync({"output_path": str(tmp_path / "c.md")}, _Ctx())
